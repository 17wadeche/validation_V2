from __future__ import annotations
import json
import re
from io import BytesIO
from typing import Iterable, Optional
try:  # pragma: no cover - optional dependency
    from docx.enum.text import WD_BREAK  # type: ignore
except Exception:  # pragma: no cover - optional dependency fallback
    WD_BREAK = None
class DocxExportError(RuntimeError):
    """Raised when a draft cannot be exported to DOCX."""
def _iter_paragraphs(document, include_headers_footers: bool = False) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell, include_headers_footers=include_headers_footers)
    if not include_headers_footers:
        return
    for section in getattr(document, "sections", []) or []:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr:
                yield from _iter_paragraphs(hdr, include_headers_footers=include_headers_footers)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr:
                yield from _iter_paragraphs(ftr, include_headers_footers=include_headers_footers)
KNOWN_PLACEHOLDER_SNIPPETS = [
    "assurance (standard deliverables) sample purpose statement",
    "validation (enhanced deliverables) sample purpose statement",
    "the purpose of this document is to record the quality assurance activities",
]
GUIDANCE_SNIPPETS = [
    "blue text is included for reference and needs to be converted to black text or removed prior to routing the document. some sections are all blue, these sections are recommended, but not required. important: the header area in this document must be left blank. mrcs d2 will automatically insert the header. this form is intended for tools that are assessed to have low or moderate risk level as assessed per quality assurance for quality data science and analytics tools (d00483500) procedure. for high risk level tools, follow the mscm procedure (document 117376-sop). for projects that are determined to be no risk per d00483500, this form is optional.",
    "blue text is included for reference and needs to be converted to black text or removed prior to routing the document.",
    "important:  the header area in this document must be left blank.",
    "this form is intended for tools that are assessed to have low or moderate risk level as assessed per quality assurance",
    "some sections are all blue, these sections are recommended, but not required.",
    "the header area in this document must be left blank.",
    "this form is intended for tools that are assessed to have low or moderate risk level",
]
def _matches_known_placeholder(text: str) -> bool:
    lowered = text.strip().lower()
    if not lowered:
        return False
    return any(snippet in lowered for snippet in KNOWN_PLACEHOLDER_SNIPPETS)
def _normalize(text: str) -> str:
    return " ".join((text or "").strip().lower().split())
def _looks_like_placeholder(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if any(stripped.startswith(prefix) for prefix in ("<", "[[", "{")):
        return True
    if any(marker in stripped for marker in ("<", ">", "[[", "]]", "{", "}")):
        return True
    return _matches_known_placeholder(stripped)
def _is_guidance_text(text: str) -> bool:
    normalized = _normalize(text)
    if not normalized:
        return False
    if normalized.startswith("blue text is included for reference"):
        return True
    return any(snippet in normalized for snippet in GUIDANCE_SNIPPETS)
def _color_matches_blue(color) -> bool:
    if not color:
        return False
    known_blues = {
        "0000ff",
        "1f4e79",
        "2f5496",
        "2b579a",
        "4472c4",  # Word Accent 1
        "5b9bd5",  # Word Accent 1 variant
        "0070c0",  # Word Accent 1 dark
        "0563c1",  # Alternate theme blue
    }
    rgb = getattr(color, "rgb", None)
    if rgb:
        rgb_text = str(rgb).lower()
        if rgb_text in known_blues:
            return True
    val = getattr(color, "val", None)
    if val and str(val).lower() in known_blues:
        return True
    try:  # pragma: no cover - depends on docx internals
        from docx.oxml.ns import qn  # type: ignore
        element = getattr(color, "_element", None)
        if element is not None:
            raw_val = element.get(qn("w:val"))
            if raw_val and raw_val.lower() in known_blues:
                return True

            theme_val = element.get(qn("w:themeColor"))
            if theme_val and "accent" in theme_val.lower():
                return True
    except Exception:
        pass
    theme_color = getattr(color, "theme_color", None)
    if theme_color:
        theme_text = str(theme_color).lower()
        if "accent" in theme_text or "blue" in theme_text:
            return True
    highlight = getattr(color, "highlight_color", None)
    if highlight:
        highlight_text = str(highlight).lower()
        if "blue" in highlight_text or "accent" in highlight_text:
            return True
    return False
def _is_blue_run(run) -> bool:
    if _color_matches_blue(getattr(run.font, "color", None)):
        return True
    try:  # pragma: no cover - optional style metadata
        run_style = getattr(run, "style", None)
        if run_style and _color_matches_blue(getattr(run_style.font, "color", None)):
            return True
    except Exception:
        pass
    try:  # pragma: no cover - paragraph styles may carry the placeholder color
        paragraph = getattr(run, "paragraph", None)
        if paragraph and _color_matches_blue(getattr(paragraph.style.font, "color", None)):
            return True
    except Exception:
        pass
    return False
def _is_placeholder_run(run) -> bool:
    text = getattr(run, "text", "") or ""
    stripped = text.strip()
    if not stripped:
        return False
    if _is_blue_run(run):
        return True
    style = getattr(run, "style", None)
    try:  # pragma: no cover - style lookups depend on template metadata
        style_name = getattr(style, "name", "") or ""
        if "placeholder" in style_name.lower():
            return True
    except Exception:
        pass
    lowered = stripped.lower()
    if any(marker in lowered for marker in ("fill", "replace", "insert", "enter", "provided by")):
        return True
    if any(token in stripped for token in ("<", ">", "[[", "]]", "{", "}", "___")):
        return True
    if _matches_known_placeholder(stripped):
        return True
    if stripped.isupper() and len(stripped) > 6:
        return True
    return False
def _paragraph_placeholder_score(paragraph) -> int:
    score = 0
    try:
        if _color_matches_blue(getattr(paragraph.style.font, "color", None)):
            score += 2
    except Exception:
        pass
    runs = getattr(paragraph, "runs", []) or []
    if not runs:
        return score
    for run in runs:
        if _is_placeholder_run(run):
            score += 3
        elif _is_blue_run(run):
            score += 2
    text = (paragraph.text or "").strip()
    if text and _looks_instructional(text):
        score += 1
    if _looks_like_placeholder(text):
        score += 2
    if _matches_known_placeholder(text):
        score += 3
    return score
def _looks_instructional(text: str) -> bool:
    lowered = text.strip().lower()
    if not lowered:
        return False
    phrases = (
        "enter ",
        "describe ",
        "provide ",
        "summarize ",
        "explain ",
        "insert ",
        "complete ",
    )
    return any(lowered.startswith(p) or f" {p}" in lowered for p in phrases)
def _replace_paragraph_with_lines(paragraph, lines: list[str]):
    style = paragraph.style
    parent = paragraph._parent
    paragraph.text = lines[0]
    paragraph.style = style
    for line in lines[1:]:
        parent.add_paragraph(line, style=style)
def _replace_run_with_draft(run, draft: str):
    run.text = ""
    parts = draft.splitlines() or [draft]
    for idx, part in enumerate(parts):
        if idx and WD_BREAK:
            try:  # pragma: no cover - relies on optional enum
                run.add_break(WD_BREAK.LINE)
            except Exception:
                run.add_text("\n")
        elif idx:
            run.add_text("\n")
        run.add_text(part)
def _extract_purpose_text(draft: str) -> Optional[str]:
    if not draft:
        return None
    paragraphs = [para.strip() for para in draft.split("\n\n") if para and para.strip()]
    for para in paragraphs:
        if "purpose of this document is" in _normalize(para):
            return para
    return paragraphs[0] if paragraphs else None
def _extract_scope_values(draft: str) -> dict[str, str]:
    if not draft:
        return {}
    patterns: list[tuple[str, str]] = [
        ("<Tool Name>", r"Tool Name:\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("<tool name>", r"Tool Name:\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("<#.#.#>", r"Tool Release:\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("<Tool Type>", r"Tool Type:\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("Tool Type", r"Tool Type:\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("<Tool Location>", r"Tool Location:\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("Tool Location", r"Tool Location:\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("<Tool Platform>", r"Tool Platform\s*\n\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("Tool Platform", r"Tool Platform\s*\n\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("<Object Type>", r"Object Type\s*\n\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("Object Type", r"Object Type\s*\n\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("<Data Source(s)>", r"Data Source\(s\):\s*([\s\S]*?)(?:\n\s*\n|$)"),
        ("Data Source(s)", r"Data Source\(s\):\s*([\s\S]*?)(?:\n\s*\n|$)"),
    ]
    extracted: dict[str, str] = {}
    for token, pattern in patterns:
        match = re.search(pattern, draft, flags=re.IGNORECASE)
        if not match:
            continue
        value = match.group(1).strip()
        if not value:
            continue
        cleaned = re.sub(r"\s+", " ", value)
        extracted.setdefault(token, cleaned)
    if "<Tool Name>" in extracted:
        extracted.setdefault("<tool name>", extracted["<Tool Name>"])
    if "<tool name>" in extracted:
        extracted.setdefault("<Tool Name>", extracted["<tool name>"])
    return extracted
def _filter_placeholder_map(raw_map: dict[str, str]) -> dict[str, str]:
    filtered: dict[str, str] = {}
    for key, value in raw_map.items():
        if not key:
            continue
        if any(marker in key for marker in ("<", ">", "[", "]", "{", "}", "___")):
            filtered[key] = value
            continue
        if _looks_like_placeholder(key):
            filtered[key] = value
    return filtered
def _replace_tokens_in_run(run, token_map: dict[str, str], token_pattern) -> bool:
    text = getattr(run, "text", "") or ""
    if not text or not token_pattern:
        return False
    def _render(match):
        token = match.group(0)
        replacement = token_map.get(token, token)
        return replacement
    replaced = token_pattern.sub(_render, text)
    if replaced != text:
        run.text = replaced
        return True
    return False
def _build_token_pattern(token_map: dict[str, str]):
    if not token_map:
        return None
    safe_tokens = [re.escape(token) for token in sorted(token_map.keys(), key=len, reverse=True)]
    if not safe_tokens:
        return None
    return re.compile("(" + "|".join(safe_tokens) + ")")
def _parse_structured_draft(draft: str) -> tuple[str, dict[str, str]]:
    if not draft:
        return "", {}
    try:
        data = json.loads(draft)
    except Exception:
        return draft.strip(), {}
    if isinstance(data, dict):
        placeholders: dict[str, str] = {}
        answers = data.get("answers")
        if isinstance(data.get("placeholders"), dict):
            placeholders = {str(k): str(v) for k, v in data["placeholders"].items()}
        elif all(isinstance(v, (str, int, float)) for v in data.values()):
            placeholders = {str(k): str(v) for k, v in data.items()}
        if not placeholders and isinstance(answers, list):
            for item in answers:
                if not isinstance(item, dict):
                    continue
                token = (
                    item.get("token")
                    or item.get("placeholder")
                    or item.get("template_text")
                    or item.get("text")
                )
                value = (
                    item.get("replacement")
                    or item.get("answer")
                    or item.get("value")
                )
                if token is None or value is None:
                    continue
                placeholders[str(token)] = str(value)
        full_text = data.get("draft") or data.get("full_text") or data.get("text")
        draft_text = str(full_text).strip() if isinstance(full_text, (str, int, float)) else ""
        return draft_text, placeholders
    return draft.strip(), {}
def draft_to_docx_bytes(draft: str, template_bytes: Optional[bytes] = None) -> bytes:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise DocxExportError(
            "python-docx is required to export Word files. Install with `pip install python-docx`."
        ) from exc
    if template_bytes:
        try:
            doc = Document(BytesIO(template_bytes))
        except Exception as exc:
            raise DocxExportError(
                "Uploaded template is not a valid .docx file. Upload a Word template or leave the template blank to export a generated draft."
            ) from exc
    else:
        doc = Document()
    structured_draft, placeholder_map = _parse_structured_draft(draft)
    extracted_scope = _extract_scope_values(structured_draft or draft)
    for key, value in extracted_scope.items():
        placeholder_map.setdefault(key, value)
    placeholder_map = _filter_placeholder_map(placeholder_map)
    placeholders = ["[[GENERATED_DRAFT]]", "<GENERATED_DRAFT>", "{GENERATED_DRAFT}"]
    token_pattern = _build_token_pattern(placeholder_map)
    generic_placeholder_pattern = _build_token_pattern({placeholder: "" for placeholder in placeholders})
    inserted = False
    map_replaced = False
    purpose_text = _extract_purpose_text(structured_draft or draft)
    replaced_purpose = False
    use_full_draft = not placeholder_map
    lines = (structured_draft or draft).splitlines() or [structured_draft or draft]
    candidate_paragraph = None
    best_scored_paragraph = None
    best_score = 0
    for paragraph in _iter_paragraphs(doc):
        runs = list(getattr(paragraph, "runs", []))
        score = _paragraph_placeholder_score(paragraph)
        if score > best_score:
            best_score = score
            best_scored_paragraph = paragraph
        normalized_text = _normalize(paragraph.text)
        if _is_guidance_text(normalized_text):
            paragraph.text = ""
            continue
        if purpose_text and (
            _matches_known_placeholder(paragraph.text or "")
            or "assurance (standard deliverables) sample purpose statement" in normalized_text
            or "validation (enhanced deliverables) sample purpose statement" in normalized_text
        ):
            if not replaced_purpose:
                _replace_paragraph_with_lines(paragraph, purpose_text.splitlines() or [purpose_text])
                replaced_purpose = True
            else:
                paragraph.text = ""
            map_replaced = True
            continue
        if replaced_purpose and (
            "assurance (standard deliverables) sample purpose statement" in normalized_text
            or "validation (enhanced deliverables) sample purpose statement" in normalized_text
        ):
            paragraph.text = ""
            continue
        para_text = (paragraph.text or "").strip()
        if placeholder_map and para_text in placeholder_map and (score or _looks_like_placeholder(para_text)):
            replacement_lines = str(placeholder_map[para_text]).splitlines() or [""]
            _replace_paragraph_with_lines(paragraph, replacement_lines)
            map_replaced = True
            continue
        if placeholder_map and token_pattern and (score or _looks_like_placeholder(paragraph.text)):
            replaced_here = False
            for run in runs:
                if _replace_tokens_in_run(run, placeholder_map, token_pattern):
                    replaced_here = True
            if not replaced_here:
                combined = paragraph.text or ""
                if token_pattern.search(combined):
                    new_text = token_pattern.sub(lambda match: placeholder_map.get(match.group(0), match.group(0)), combined)
                    _replace_paragraph_with_lines(paragraph, new_text.splitlines() or [new_text])
                    replaced_here = True
            map_replaced = map_replaced or replaced_here
            if replaced_here:
                continue
        if use_full_draft and _matches_known_placeholder(paragraph.text or ""):
            _replace_paragraph_with_lines(paragraph, lines)
            inserted = True
            continue
        for idx, run in enumerate(runs):
            if _replace_tokens_in_run(run, {placeholder: "\n".join(lines) for placeholder in placeholders}, generic_placeholder_pattern):
                inserted = True
                break
            if use_full_draft and _is_placeholder_run(run):
                _replace_run_with_draft(run, structured_draft or draft)
                for follower in runs[idx + 1 :]:
                    if _is_placeholder_run(follower):
                        follower.text = ""
                    else:
                        break
                inserted = True
                break
        if inserted:
            continue
        if use_full_draft and any(_looks_like_placeholder(getattr(run, "text", "")) for run in runs):
            _replace_paragraph_with_lines(paragraph, lines)
            inserted = True
            continue
        if use_full_draft and _looks_like_placeholder(paragraph.text):
            _replace_paragraph_with_lines(paragraph, lines)
            inserted = True
            continue
        if not candidate_paragraph and _looks_instructional(paragraph.text):
            candidate_paragraph = paragraph
        for placeholder in placeholders:
            if placeholder in paragraph.text:
                if use_full_draft:
                    _replace_paragraph_with_lines(paragraph, lines)
                else:
                    paragraph.text = paragraph.text.replace(placeholder, "")
                inserted = True
                break
    if not inserted and not map_replaced and candidate_paragraph:
        _replace_paragraph_with_lines(candidate_paragraph, lines)
        inserted = True
    if use_full_draft and not inserted and not map_replaced and best_scored_paragraph and best_score >= 3:
        _replace_paragraph_with_lines(best_scored_paragraph, lines)
        inserted = True
    if not inserted and not map_replaced:
        if use_full_draft:
            for line in lines:
                doc.add_paragraph(line)
        elif structured_draft or draft:
            for line in lines:
                doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
def docx_bytes_to_html(docx_bytes: bytes) -> Optional[str]:
    try:  # pragma: no cover - optional dependency
        import mammoth  # type: ignore
    except Exception:
        return None
    try:
        result = mammoth.convert_to_html(BytesIO(docx_bytes), style_map="p[style-name='Normal'] => p")
    except Exception:
        return None
    return result.value if result else None