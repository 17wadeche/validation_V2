from __future__ import annotations
import json
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple
from io import BytesIO
from datetime import datetime
import sys
import subprocess
from docx import Document
import re
from typing import Any
from docx.shared import Pt
from flask import Flask, render_template_string, request, send_file
from src.validation_agent import __version__ as APP_VERSION
from src.validation_agent.prompt_builder import (
    Example,
    build_planning_prompt,
    build_prompt,
    build_update_prompt,
    extract_placeholders,
    build_design_update_prompt,
    build_functional_requirements_prompt,
    build_testing_alignment_prompt,
)
from src.validation_agent.telemetry import log_event, Timer
from src.validation_agent.document_loader import load_text_document
from src.validation_agent.medtronic_client import MedtronicGPTClient, MedtronicGPTError
from src.validation_agent.credentials import StoredCredentials, load_credentials, save_credentials
from src.validation_agent.storage import (
    SavedInputs,
    StoredFile,
    load_saved_inputs,
    save_inputs,
)
from src.validation_agent.workbook_loader import extract_excel_context, extract_pbix_context
import tempfile
import logging
import time
import uuid
from contextlib import contextmanager
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
app = Flask(__name__)
@contextmanager
def _timed(timings: dict, name: str):
    start = time.perf_counter()
    try:
        yield
    finally:
        timings[name] = round((time.perf_counter() - start) * 1000, 2)
def _read_upload(file_storage) -> Tuple[Optional[str], Optional[bytes], Optional[str], Optional[str]]:
    if not file_storage:
        return None, None, None, None
    filename = file_storage.filename
    if not filename:
        return None, None, None, None
    raw_bytes = b""
    try:
        file_storage.stream.seek(0)
        raw_bytes = file_storage.stream.read()
    except Exception:
        try:
            raw_bytes = file_storage.read()
        except Exception:
            raw_bytes = b""
    if raw_bytes is None:
        raw_bytes = b""
    suffix = Path(filename).suffix
    text: Optional[str]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(raw_bytes)
        tmp.flush()
        try:
            text = load_text_document(Path(tmp.name))
        except Exception:
            text = ""
    return text, raw_bytes, suffix, filename
def _read_saved_file(saved_file: StoredFile) -> Tuple[Optional[str], Optional[bytes], Optional[str], Optional[str]]:
    try:
        raw_bytes = saved_file.to_bytes()
    except Exception:
        return None, None, None, None
    suffix = saved_file.suffix or Path(saved_file.name).suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(raw_bytes)
        tmp.flush()
        text = load_text_document(Path(tmp.name))
    return text, raw_bytes, suffix, saved_file.name
def _dedupe_by_name(files: List[StoredFile]) -> List[StoredFile]:
    seen = set()
    unique: List[StoredFile] = []
    for item in files:
        if item.name in seen:
            continue
        seen.add(item.name)
        unique.append(item)
    return unique
def _extract_questions_from_json(payload: str) -> List[str]:
    if not payload:
        return []
    try:
        data = json.loads(payload)
    except Exception:
        return []
    questions: List[str] = []
    if isinstance(data, dict):
        for key in ("questions", "clarifying_questions"):
            value = data.get(key)
            if isinstance(value, list):
                questions.extend(str(item).strip() for item in value if str(item).strip())
    return questions
def _gather_examples(
    uploaded_files,
    saved_examples: List[StoredFile],
    *,
    tag: Optional[str] = None,
) -> Tuple[List[Example], List[StoredFile]]:
    examples: List[Example] = []
    stored_examples: List[StoredFile] = []
    for file_storage in uploaded_files or []:
        content, raw_bytes, _, filename = _read_upload(file_storage)
        if raw_bytes is not None and filename:
            label = f"[{tag}] {filename}" if tag else filename
            examples.append(Example(title=label, context="", output=content or ""))
            stored_examples.append(StoredFile.from_bytes(label, raw_bytes))
    for saved in saved_examples:
        content, raw_bytes, _, name = _read_saved_file(saved)
        if raw_bytes is not None and name:
            examples.append(Example(title=name, context="", output=content or ""))
            stored_examples.append(StoredFile.from_bytes(name, raw_bytes))
    return examples, _dedupe_by_name(stored_examples)
def _format_code_section(label: str, snippets: List[str]) -> str:
    if not snippets:
        return ""
    return f"## {label}\n" + "\n".join(snippets)
def _gather_code_context(
    current_code_files,
    inline_code: str,
    old_code_files=None,
    new_code_files=None,
    inline_code_old: str = "",
    inline_code_new: str = "",
) -> str:
    snippets_current: List[str] = []
    snippets_old: List[str] = []
    snippets_new: List[str] = []
    def _extract_from_upload(fs) -> str:
        from pathlib import Path
        if not fs or not fs.filename:
            return ""
        suffix = Path(fs.filename).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            fs.stream.seek(0)
            tmp.write(fs.stream.read())
            tmp.flush()
            tmp_path = Path(tmp.name)
        try:
            if suffix in {".xlsm", ".xlsx", ".xls", ".xlsb"}:
                return extract_excel_context(tmp_path)
            if suffix == ".pbix":
                return extract_pbix_context(tmp_path)
            tmp_path_str = tmp_path.read_text(encoding="utf-8", errors="ignore")
            return tmp_path_str
        except Exception:
            return ""
    def _collect(files, bucket: List[str], tag: str):
        from pathlib import Path
        for fs in files or []:
            if not fs or not fs.filename:
                continue
            text = _extract_from_upload(fs)
            if text.strip():
                bucket.append(f"\n# {tag} File: {fs.filename}\n{text.strip()}\n")
    _collect(current_code_files, snippets_current, "Current")
    _collect(old_code_files or [], snippets_old, "Previous")
    _collect(new_code_files or [], snippets_new, "Updated")
    sections = [
        _format_code_section("Current code/context", snippets_current),
        _format_code_section("Previous version (for updates)", snippets_old),
        _format_code_section("Updated version (for updates)", snippets_new),
    ]
    if inline_code.strip():
        sections.append(
            _format_code_section(
                "Additional notes (current/general)", [inline_code.strip()]
            )
        )
    if inline_code_old.strip():
        sections.append(
            _format_code_section("Additional notes (previous)", [inline_code_old.strip()])
        )
    if inline_code_new.strip():
        sections.append(
            _format_code_section("Additional notes (updated)", [inline_code_new.strip()])
        )
    return "\n\n".join(part for part in sections if part).strip()
def _apply_functional_requirements_enrichment(
    *,
    draft: str,
    template_text: str,
    prompt: str | None,
    code_context: str,
    plan_context: str,
    client: MedtronicGPTClient,
    model: str,
) -> tuple[str, Optional[str]]:
    if not draft or not template_text:
        return draft, None
    try:
        parsed = json.loads(draft)
    except json.JSONDecodeError:
        return draft, None
    if not isinstance(parsed, dict):
        return draft, None
    placeholders_map = parsed.get("placeholders")
    if not isinstance(placeholders_map, dict):
        placeholders_map = {}
    tokens = extract_placeholders(template_text)
    fr_token: Optional[str] = None
    import re
    def _norm(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", s.lower())
    candidate_keys: list[str] = []
    if isinstance(placeholders_map, dict):
        candidate_keys.extend(placeholders_map.keys())
    if isinstance(tokens, (list, tuple)):
        candidate_keys.extend(tokens)
    for key in candidate_keys:
        n = _norm(str(key))
        if "functionalrequirements" in n:
            fr_token = key
            break
    if fr_token is None:
        for key in candidate_keys:
            n = _norm(str(key))
            if n == "requirements" or n.endswith("requirements"):
                fr_token = key
                break
    fr_prompt = build_functional_requirements_prompt(
        prompt or "",
        code_context,
        plan_context=plan_context,
    )
    try:
        fr_raw = client.generate_completion(fr_prompt, model=model)
    except MedtronicGPTError as exc:
        logging.getLogger(__name__).warning(
            "Functional requirements enrichment failed: %s", exc
        )
        return draft, None
    try:
        fr_data = json.loads(fr_raw)
    except json.JSONDecodeError:
        return draft, fr_raw
    if isinstance(fr_data, dict) and "functional_requirements" in fr_data:
        fr_list = fr_data["functional_requirements"]
    else:
        fr_list = fr_data
    functional_reqs: list[dict] = []
    if isinstance(fr_list, list):
        version_value = ""
        version_source = parsed.get("placeholders")
        if isinstance(version_source, dict):
            for key in ("<#.#.#>", "<Tool Release>", "<Release Version>"):
                v = version_source.get(key)
                if isinstance(v, str) and v.strip():
                    version_value = v.strip()
                    break
        for idx, item in enumerate(fr_list, start=1):
            if isinstance(item, dict):
                desc = item.get("Description") or item.get("description") or ""
                req_id = item.get("Unique Req ID") or item.get("id") or ""
            else:
                desc = str(item)
                req_id = ""
            desc = str(desc).strip()
            if not desc:
                continue
            if not isinstance(req_id, str) or not req_id.strip():
                req_id = f"F{idx}"
            functional_reqs.append(
                {
                    "Unique Req ID": req_id,
                    "Description": desc,
                    "Release Implemented": version_value or "1.0",
                }
            )
    if not functional_reqs:
        return draft, fr_raw
    if not fr_token:
        fr_token = "<Functional Requirements>"
    placeholders_map[fr_token] = functional_reqs
    parsed["placeholders"] = placeholders_map
    updated = json.dumps(parsed, indent=2)
    return updated, fr_raw
def _apply_testing_documentation_alignment_enrichment(
    *,
    draft: str,
    template_text: str,
    client: MedtronicGPTClient,
    model: str,
) -> tuple[str, Optional[str]]:
    if not draft:
        return draft, None
    try:
        parsed = json.loads(draft)
    except Exception:
        return draft, None
    if not isinstance(parsed, dict):
        return draft, None
    placeholders_map = parsed.get("placeholders")
    if not isinstance(placeholders_map, dict):
        placeholders_map = {}
    answers_list = parsed.get("answers")
    if not isinstance(answers_list, list):
        answers_list = []
    import re
    def _norm(s: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", s.lower())
    fr_value = None
    for k, v in placeholders_map.items():
        if "functionalrequirements" in _norm(str(k)) and isinstance(v, list):
            fr_value = v
            break
    if fr_value is None:
        fr_value = placeholders_map.get("<Functional Requirements>")
    if not isinstance(fr_value, list) or not fr_value:
        return draft, None  # nothing to align to
    test_token = None
    for k in list(placeholders_map.keys()) + extract_placeholders(template_text or ""):
        n = _norm(str(k))
        if n in {"testingdocumentation", "testingdoc", "testdocumentation"}:
            test_token = str(k)
            break
    existing_testing = None
    if test_token and test_token in placeholders_map:
        existing_testing = placeholders_map.get(test_token)
    else:
        for a in answers_list:
            if not isinstance(a, dict):
                continue
            ph = (a.get("placeholder") or a.get("token") or "").strip()
            if "testingdocumentation" in _norm(ph):
                test_token = ph
                existing_testing = a.get("replacement", a.get("answer"))
                break
    if test_token is None or existing_testing is None:
        return draft, None  # no testing section present
    align_prompt = build_testing_alignment_prompt(
        functional_requirements=fr_value,
        existing_testing_doc=existing_testing,
    )
    try:
        raw = client.generate_completion(align_prompt, model=model)
    except MedtronicGPTError:
        return draft, None
    try:
        payload = json.loads(_extract_json_from_reply(raw) or raw)
    except Exception:
        return draft, raw
    aligned_value = None
    if isinstance(payload, dict):
        if isinstance(payload.get("testing_documentation"), list):
            aligned_value = payload["testing_documentation"]
        elif isinstance(payload.get("testing_documentation_text"), str):
            aligned_value = payload["testing_documentation_text"].strip()
    else:
        aligned_value = payload
    if aligned_value is None:
        return draft, raw
    if test_token in placeholders_map:
        placeholders_map[test_token] = aligned_value
    else:
        for a in answers_list:
            if isinstance(a, dict) and (a.get("placeholder") == test_token or a.get("token") == test_token):
                a["replacement"] = aligned_value
                break
    parsed["placeholders"] = placeholders_map
    parsed["answers"] = answers_list
    return json.dumps(parsed, indent=2), raw
def _extract_json_from_reply(reply: str) -> str:
    if not reply:
        return ""
    s = reply.strip()
    if s.startswith("```"):
        s = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", s)
        s = re.sub(r"\s*```$", "", s).strip()
    first_obj = s.find("{")
    first_arr = s.find("[")
    starts = [i for i in (first_obj, first_arr) if i != -1]
    if starts:
        start = min(starts)
        end_obj = s.rfind("}")
        end_arr = s.rfind("]")
        end = max(end_obj, end_arr)
        if end > start:
            s = s[start : end + 1].strip()
    return s
def _compute_missing_placeholders(template_text: str, draft_json: str) -> List[str]:
    if not template_text or not draft_json:
        return []
    tokens = [tok.strip() for tok in extract_placeholders(template_text) if tok.strip()]
    if not tokens:
        return []
    try:
        data = json.loads(draft_json)
    except Exception:
        return tokens
    filled = set()
    placeholders_map = data.get("placeholders") if isinstance(data, dict) else {}
    if isinstance(placeholders_map, dict):
        for token, value in placeholders_map.items():
            token = str(token).strip()
            if not token:
                continue
            if isinstance(value, str):
                if value.strip():
                    filled.add(token)
            elif value is not None:
                filled.add(token)
    answers = data.get("answers") if isinstance(data, dict) else []
    if isinstance(answers, list):
        for entry in answers:
            if not isinstance(entry, dict):
                continue
            token = str(entry.get("placeholder", "")).strip()
            replacement = entry.get("replacement", entry.get("answer"))
            if not token:
                continue
            if replacement is None:
                continue
            if isinstance(replacement, str) and not replacement.strip():
                continue
            filled.add(token)
    return [tok for tok in tokens if tok not in filled]
def _build_prompt_from_request(
    form,
    files,
    selected_template: Optional[StoredFile],
    kept_saved_examples: List[StoredFile],
    user_instructions: str | None,
    plan_context: str | None,
    release_type: str,
) -> Tuple[str, Optional[bytes], Optional[StoredFile], List[StoredFile], str, List[Example], str]:
    template_text, template_bytes, _, template_name = _read_upload(files.get("template_file"))
    stored_template: Optional[StoredFile] = None
    if not template_text and selected_template:
        template_text, template_bytes, _, template_name = _read_saved_file(selected_template)
    if template_bytes is not None and template_name:
        stored_template = StoredFile.from_bytes(template_name, template_bytes)
        selected_template = stored_template
        selected_template_name = stored_template.name
    examples, stored_examples = _gather_examples(files.getlist("examples"), kept_saved_examples)
    code_context = _gather_code_context(
        files.getlist("code_files"),
        form.get("code_context", ""),
        old_code_files=files.getlist("code_files_old") if release_type == "update" else None,
        new_code_files=files.getlist("code_files_new") if release_type == "update" else None,
        inline_code_old=form.get("code_context_old", ""),
        inline_code_new=form.get("code_context_new", ""),
    )
    prompt = build_prompt(
        template_text or "",
        examples,
        code_context,
        plan_context=plan_context,
        release_type=release_type,
        user_instructions=user_instructions,
    )
    return (
        prompt,
        template_bytes,
        stored_template,
        stored_examples,
        template_text or "",
        examples,
        code_context,
    )
def _open_docx_file(path: str) -> None:
    try:
        if sys.platform.startswith("win"):
            import os
            os.startfile(path)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])
        else:
            subprocess.Popen(["xdg-open", path])
    except Exception as exc:
        app.logger.warning("Could not auto-open docx: %s", exc)
def _order_keys_by_template(keys: list[str], template_text: str) -> list[str]:
    if not template_text:
        return list(keys)
    scored = []
    for idx, k in enumerate(keys):
        p = template_text.find(k)
        scored.append((p if p != -1 else 10**18, idx, k))
    scored.sort(key=lambda t: (t[0], t[1]))
    return [k for _, _, k in scored]
def _add_value_to_doc(doc: Document, value):
    if value is None or value == "":
        doc.add_paragraph("(empty)")
        return
    if isinstance(value, str):
        doc.add_paragraph(value)
        return
    if isinstance(value, list):
        if not value:
            doc.add_paragraph("(empty list)")
            return
        for item in value:
            if isinstance(item, dict):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("")  # anchor
                for k, v in item.items():
                    sp = doc.add_paragraph(style="List Bullet 2")
                    sp.add_run(f"{k}: ").bold = True
                    sp.add_run(str(v))
            else:
                doc.add_paragraph(str(item), style="List Bullet")
        return
    if isinstance(value, dict):
        for k, v in value.items():
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{k}: ").bold = True
            p.add_run(str(v))
        return
    doc.add_paragraph(str(value))
def _build_docx_bytes_for_view(
    *,
    view: str,
    draft_json: str,
    template_text: str = "",
    markdown_text: str = "",
) -> bytes:
    doc = Document()
    doc.add_heading("Validation Export", level=1)
    doc.add_paragraph(f"View: {view}")
    if view == "json":
        doc.add_heading("JSON", level=2)
        for line in (draft_json or "").splitlines():
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
    elif view == "markdown":
        doc.add_heading("Markdown", level=2)
        for line in (markdown_text or "").splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_heading("Answers", level=2)
        try:
            parsed = json.loads(draft_json) if draft_json else {}
        except Exception:
            parsed = {}
        placeholders = parsed.get("placeholders") if isinstance(parsed, dict) else {}
        if not isinstance(placeholders, dict):
            placeholders = {}
        answers_list = parsed.get("answers") if isinstance(parsed, dict) else []
        if not isinstance(answers_list, list):
            answers_list = []
        questions = parsed.get("questions") if isinstance(parsed, dict) else []
        if not isinstance(questions, list):
            questions = []
        key_list: list[str] = []
        seen = set()
        for k in placeholders.keys():
            if k not in seen:
                key_list.append(k)
                seen.add(k)
        for item in answers_list:
            if isinstance(item, dict):
                ph = item.get("placeholder") or item.get("token")
                if isinstance(ph, str):
                    ph = ph.strip()
                    if ph and ph not in seen:
                        key_list.append(ph)
                        seen.add(ph)
        ordered = _order_keys_by_template(key_list, template_text or "")
        for ph in ordered:
            val = placeholders.get(ph)
            p = doc.add_paragraph()
            p.add_run(ph).bold = True
            if val is None or val == "":
                found = None
                for a in answers_list:
                    if isinstance(a, dict) and (a.get("placeholder") == ph or a.get("token") == ph):
                        repl = a.get("replacement", a.get("answer"))
                        if repl not in (None, ""):
                            found = repl
                            break
                val = found
            if val is None or val == "":
                continue  
            _add_value_to_doc(doc, val)
            doc.add_paragraph("")  # spacer
        if questions:
            doc.add_heading("Remaining questions", level=2)
            for q in questions:
                if isinstance(q, str) and q.strip():
                    doc.add_paragraph(q.strip(), style="List Bullet")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
@app.route("/export_docx", methods=["POST"])
def export_docx():
    view = (request.form.get("view") or "friendly").strip().lower()
    draft_json = request.form.get("draft_json", "") or ""
    template_text = request.form.get("template_text", "") or ""
    markdown_text = request.form.get("markdown_text", "") or ""
    docx_bytes = _build_docx_bytes_for_view(
        view=view,
        draft_json=draft_json,
        template_text=template_text,
        markdown_text=markdown_text,
    )
    tmp_path = Path(tempfile.gettempdir()) / f"validation_export_{view}_{uuid.uuid4().hex}.docx"
    try:
        tmp_path.write_bytes(docx_bytes)
        if request.form.get("open_now") == "1":
            _open_docx_file(str(tmp_path))
    except Exception as exc:
        app.logger.warning("Export temp write/open failed: %s", exc)
    filename = f"validation_export_{view}.docx"
    return send_file(
        BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )
@app.route("/", methods=["GET", "POST"])
def index():
    app.logger.info("index() called, method=%s", request.method)
    prompt: Optional[str] = None
    draft: Optional[str] = None
    error: Optional[str] = None
    history: list[dict] = []
    plan_text: str = ""
    draft_questions: List[str] = []
    template_bytes: Optional[bytes] = None
    code_context_text: str = ""
    code_context_old_text: str = ""
    code_context_new_text: str = ""
    template_text: str = ""
    user_instructions = request.form.get("user_instructions", "") or ""
    missing_placeholders: List[str] = []
    coverage_note: Optional[str] = None
    stored_inputs: SavedInputs = load_saved_inputs()
    persisted_inputs: SavedInputs = stored_inputs
    draft_json_from_form: str = ""
    release_type: str = "initial"
    design_text: str = ""
    fr_raw: str = "" 
    selected_template_name: str = stored_inputs.templates[0].name if stored_inputs.templates else ""
    defaults = {
        "base_url": MedtronicGPTClient.DEFAULT_BASE_URL,
        "api_version": MedtronicGPTClient.DEFAULT_API_VERSION,
        "path_template": MedtronicGPTClient.DEFAULT_PATH_TEMPLATE,
        "model": "gpt-41",
    }
    stored = load_credentials()
    defaults.update(
        {
            "base_url": stored.base_url or defaults["base_url"],
            "api_version": stored.api_version or defaults["api_version"],
            "path_template": stored.path_template or defaults["path_template"],
            "model": getattr(stored, "model", "") or defaults["model"],
        }
    )
    if request.method == "POST":
        run_id = uuid.uuid4().hex
        timings: dict[str, float] = {}
        run_started_ms = round(time.time() * 1000)
        run_action = request.form.get("action", "build")
        run_error: Optional[str] = None
        run_success = False
        draft_json_from_form = request.form.get("draft_json", "")
        remove_template_name = request.form.get("remove_template", "").strip()
        remove_example_name = request.form.get("remove_example", "").strip()
        clear_saved_templates = request.form.get("clear_templates") == "on"
        keep_saved_templates = not clear_saved_templates
        kept_saved_examples: List[StoredFile] = []
        for idx, saved_example in enumerate(stored_inputs.examples):
            keep_flag = request.form.get(f"keep_example_{idx}")
            if keep_flag == "on":
                kept_saved_examples.append(saved_example)
        template_choice = request.form.get("selected_template", "")
        selected_template_name = template_choice.strip()
        available_templates = stored_inputs.templates if keep_saved_templates else []
        selected_template_file: Optional[StoredFile] = None
        if available_templates and selected_template_name:
            selected_template_file = next(
                (item for item in available_templates if item.name == selected_template_name),
                available_templates[0],
            )
            selected_template_name = selected_template_file.name
        final_templates: List[StoredFile] = available_templates
        final_examples: List[StoredFile] = kept_saved_examples
        plan_text = request.form.get("plan_text", "")
        remember_credentials = request.form.get("remember_credentials") == "on"
        release_type = request.form.get("release_type", "initial") or "initial"
        history_json = request.form.get("history_json", "[]")
        code_context_text = request.form.get("code_context", "")
        code_context_old_text = request.form.get("code_context_old", "")
        code_context_new_text = request.form.get("code_context_new", "")
        try:
            history = json.loads(history_json) if history_json else []
        except json.JSONDecodeError:
            history = []
        action = request.form.get("action", "build")
        if remove_template_name:
            action = "remove_template"
        if remove_example_name:
            action = "remove_example"
        if action == "remove_template":
            updated_templates = [
                tmpl for tmpl in stored_inputs.templates if tmpl.name != remove_template_name
            ]
            persisted_inputs = SavedInputs(templates=updated_templates, examples=stored_inputs.examples)
            save_inputs(persisted_inputs)
            stored_inputs = persisted_inputs
            selected_template_name = updated_templates[0].name if updated_templates else ""
            draft = draft_json_from_form or draft
            draft_questions = _extract_questions_from_json(draft) if draft else []
            return render_template_string(
                TEMPLATE,
                prompt=prompt,
                draft=draft,
                error=error,
                defaults=defaults,
                history=history,
                stored=stored,
                saved_inputs=persisted_inputs,
                user_instructions=user_instructions,
                plan_text=plan_text,
                draft_questions=draft_questions,
                app_version=APP_VERSION,
                draft_json=draft_json_from_form,
                code_context=code_context_text,
                code_context_old=code_context_old_text,
                code_context_new=code_context_new_text,
                release_type=release_type,
                selected_template_name=selected_template_name,
                missing_placeholders=missing_placeholders,
                fr_raw=fr_raw,
            )
        if action == "remove_example":
            updated_examples = [
                ex for ex in stored_inputs.examples if ex.name != remove_example_name
            ]
            persisted_inputs = SavedInputs(templates=stored_inputs.templates, examples=updated_examples)
            save_inputs(persisted_inputs)
            stored_inputs = persisted_inputs
            kept_saved_examples = updated_examples
            draft = draft_json_from_form or draft
            draft_questions = _extract_questions_from_json(draft) if draft else []
            return render_template_string(
                TEMPLATE,
                app_version=APP_VERSION,
                prompt=prompt,
                draft=draft,
                error=error,
                defaults=defaults,
                history=history,
                stored=stored,
                saved_inputs=persisted_inputs,
                user_instructions=user_instructions,
                plan_text=plan_text,
                draft_questions=draft_questions,
                draft_json=draft_json_from_form,
                code_context=code_context_text,
                code_context_old=code_context_old_text,
                code_context_new=code_context_new_text,
                release_type=release_type,
                selected_template_name=selected_template_name,
                missing_placeholders=missing_placeholders,
                fr_raw=fr_raw,
            )
        if action == "clear":
            prompt = None
            draft = None
            error = None
            history = []
            plan_text = ""
            draft_questions = []
            draft_json_from_form = ""
            code_context_text = ""
            code_context_old_text = ""
            code_context_new_text = ""
            coverage_note = None
            template_text = ""
            missing_placeholders = []
            release_type = "initial"
            persisted_inputs = stored_inputs
            return render_template_string(
                TEMPLATE,
                prompt=prompt,
                draft=draft,
                error=error,
                app_version=APP_VERSION,
                defaults=defaults,
                history=history,
                stored=stored,
                saved_inputs=persisted_inputs,
                user_instructions=user_instructions,
                plan_text=plan_text,
                draft_questions=draft_questions,
                draft_json=draft_json_from_form,
                code_context=code_context_text,
                code_context_old=code_context_old_text,
                code_context_new=code_context_new_text,
                coverage_note=coverage_note,
                template_text=template_text,
                missing_placeholders=missing_placeholders,
                release_type=release_type,
                selected_template_name=selected_template_name,
                fr_raw=fr_raw,
          )
        if action == "feedback":
          user_name = request.form.get("user_name", "").strip()
          feedback_text = request.form.get("feedback_text", "").strip()
          feedback_draft = request.form.get("draft_json", "").strip()
          easy_view = request.form.get("easy_view", "")
          markdown_view = request.form.get("markdown_view", "")
          feedback_model = request.form.get("model", "").strip() or defaults["model"]
          try:
              log_event(
                  "ui_feedback",
                  user_name=user_name,
                  app_version=APP_VERSION,
                  action="feedback",
                  model=feedback_model,
                  success=True,
                  payload={
                      "feedback_text": feedback_text,
                      "draft_json": feedback_draft,
                      "easy_view": easy_view,
                      "markdown_view": markdown_view,
                      "release_type": request.form.get("release_type", release_type),
                  },
              )
          except Exception as e:
              app.logger.exception("Feedback telemetry failed: %s", e)
          draft = feedback_draft or draft_json_from_form or draft
          draft_json_from_form = draft or ""
          draft_questions = _extract_questions_from_json(draft) if draft else []
          return render_template_string(
              TEMPLATE,
              prompt=prompt,
              draft=draft,
              app_version=APP_VERSION,
              error=error,
              defaults=defaults,
              history=history,
              stored=stored,
              saved_inputs=persisted_inputs,
              user_instructions=user_instructions,
              plan_text=plan_text,
              draft_questions=draft_questions,
              code_context=code_context_text,
              code_context_old=code_context_old_text,
              code_context_new=code_context_new_text,
              draft_json=draft_json_from_form,
              coverage_note=coverage_note,
              template_text=template_text,
              missing_placeholders=missing_placeholders,
              release_type=release_type,
              selected_template_name=selected_template_name,
              design_text=design_text,
              fr_raw=fr_raw,
          )
        with _timed(timings, "gather_inputs_ms"):
            (
                prompt,
                template_bytes,
                stored_template,
                stored_examples,
                template_text,
                examples,
                code_context,
            ) = _build_prompt_from_request(
                request.form,
                request.files,
                selected_template_file,
                kept_saved_examples,
                user_instructions,
                plan_text,
                release_type,
            )
        if stored_template:
            selected_template_name = stored_template.name
        code_context_text = code_context
        if not template_bytes and selected_template_file:
            try:
                template_bytes = selected_template_file.to_bytes()
            except Exception:
                template_bytes = None
        client = None
        model = request.form.get("model", "").strip() or defaults["model"]
        if request.form.get("use_model") == "on":
            client = MedtronicGPTClient(
                base_url=request.form.get("base_url", "").strip() or defaults["base_url"],
                api_version=request.form.get("api_version", "").strip() or defaults["api_version"],
                path_template=request.form.get("path_template", "").strip() or defaults["path_template"],
                subscription_key=request.form.get("subscription_key", "").strip(),
                api_token=request.form.get("api_token", "").strip(),
                refresh_token=request.form.get("refresh_token", "").strip(),
            )
            if remember_credentials:
                stored = StoredCredentials(
                    subscription_key=request.form.get("subscription_key", "").strip(),
                    api_token=request.form.get("api_token", "").strip(),
                    refresh_token=request.form.get("refresh_token", "").strip(),
                    api_version=request.form.get("api_version", "").strip() or defaults["api_version"],
                    base_url=request.form.get("base_url", "").strip() or defaults["base_url"],
                    path_template=request.form.get("path_template", "").strip() or defaults["path_template"],
                    model=model,
                )
                save_credentials(stored)
        if action in {"answers", "refine"}:
            answered = []
            for key, value in request.form.items():
                if key.startswith("question_"):
                    idx = key.split("_", 1)[1]
                    question = value.strip()
                    answer = request.form.get(f"answer_{idx}", "").strip()
                    if question and answer:
                        answered.append((question, answer))
            if not draft_json_from_form.strip():
                error = "No draft JSON was provided to update with answers."
            else:
                try:
                    parsed = json.loads(draft_json_from_form)
                except json.JSONDecodeError:
                    parsed = None
                    error = "Draft JSON could not be parsed."
                if parsed is not None:
                    answers_list = parsed.get("answers")
                    if not isinstance(answers_list, list):
                        answers_list = []
                    questions_list = parsed.get("questions")
                    if not isinstance(questions_list, list):
                        questions_list = []

                    if answered:
                        for question, answer in answered:
                            answers_list.append({"question": question, "answer": answer})
                            questions_list = [q for q in questions_list if q != question]

                    parsed["answers"] = answers_list
                    parsed["questions"] = questions_list

                    if action == "answers":
                        draft = json.dumps(parsed, indent=2)
                        draft_questions = [q for q in questions_list if q]
                    elif action == "refine":
                        if not client:
                            error = "Provide MedtronicGPT credentials to update with GPT."
                        else:
                            missing_for_update: List[str] = []
                            if template_text:
                                missing_for_update = _compute_missing_placeholders(
                                    template_text, json.dumps(parsed)
                                )
                            update_prompt = build_update_prompt(
                                template_text or "",
                                examples,
                                code_context,
                                json.dumps(parsed),
                                answered,
                                plan_context=plan_text,
                                release_type=release_type,
                                missing_tokens=missing_for_update,
                                user_instructions=user_instructions,
                            )
                            try:
                                with _timed(timings, "coverage_fill_completion_ms"):
                                  draft = client.generate_completion(update_prompt, model=model)
                                draft_questions = _extract_questions_from_json(draft)
                            except MedtronicGPTError as exc:
                                error = str(exc)
        if action == "chat" and client:
            user_message = request.form.get("chat_input", "").strip()
            chat_update_json = request.form.get("chat_update_json") == "on"
            if user_message:
                history.append({"role": "user", "content": user_message})
                if chat_update_json:
                    seed = {
                        "role": "system",
                        "content": (
                            "You are assisting with Medtronic validation drafting.\n"
                            "IMPORTANT: Return VALID JSON ONLY (no markdown, no code fences).\n"
                            "You MUST output the FULL updated draft JSON in the SAME schema as the current draft:\n"
                            "{\n"
                            "  \"placeholders\": { ... },\n"
                            "  \"answers\": [ ... ],\n"
                            "  \"questions\": [ ... ],\n"
                            "  \"coverage\": { ... }\n"
                            "}\n"
                            "Start from the provided current draft JSON and apply the user's request.\n"
                            "If information is missing, add concise items to `questions` rather than guessing.\n"
                            "Never fabricate person names or signatures.\n"
                        ),
                    }
                else:
                    seed = {
                        "role": "system",
                        "content": (
                            "You are assisting with Medtronic validation drafting. Answer user questions directly and succinctly using the provided context. "
                            "Do not invent person names or signatures. If context is missing, ask one concise follow-up question. Avoid returning JSON unless explicitly requested."
                        ),
                    }
                full_history = [seed]
                if user_instructions.strip():
                  full_history.append(
                      {"role": "system", "content": "Additional user instructions:\n" + user_instructions.strip()}
                  )
                if prompt and not chat_update_json:
                  full_history.append({"role": "system", "content": f"Reference prompt context to stay on-topic:\n{prompt}"})
                if plan_text.strip():
                    full_history.append(
                        {
                            "role": "system",
                            "content": "Use this planning JSON to map placeholders before proposing edits:\n"
                            + plan_text.strip(),
                        }
                    )
                context_json = draft or draft_json_from_form
                if context_json and context_json.strip():
                    full_history.append(
                        {
                            "role": "system",
                            "content": "Ground answers in the latest generated JSON (placeholders/answers/questions):\n"
                            + context_json.strip(),
                        }
                    )
                full_history += history
                try:
                    with _timed(timings, "chat_completion_ms"):
                        reply = client.generate_completion(model=model, messages=full_history)
                    history.append({"role": "assistant", "content": reply})
                    if chat_update_json:
                      try:
                          cleaned = _extract_json_from_reply(reply) or (reply or "")
                          parsed = json.loads(cleaned)
                          if not isinstance(parsed, dict):
                              raise ValueError("Chat returned JSON but not an object.")
                          draft = json.dumps(parsed, indent=2)
                          draft_json_from_form = draft
                          draft_questions = _extract_questions_from_json(draft)
                      except Exception as e:
                          app.logger.warning("Chat JSON parse failed: %s; reply_head=%r", e, (reply or "")[:400])
                          error = "Chat JSON mode was enabled, but the model response was not valid draft JSON. Try again or disable the toggle."
                except MedtronicGPTError as exc:
                    error = str(exc)
        elif action == "chat" and not client:
            error = "Provide MedtronicGPT credentials to chat."
        elif action == "build" and client:
            if not plan_text.strip():
                planning_prompt = build_planning_prompt(
                    template_text or "", examples, code_context, release_type=release_type,
                    user_instructions=user_instructions,
                )
                try:
                    with _timed(timings, "planning_completion_ms"):
                        plan_text = client.generate_completion(planning_prompt, model=model)
                except MedtronicGPTError as exc:
                    error = str(exc)
            if not error:
                prompt = build_prompt(
                    template_text or "",
                    examples,
                    code_context,
                    plan_context=plan_text,
                    release_type=release_type,
                )
                try:
                    with _timed(timings, "build_completion_ms"):
                        draft = client.generate_completion(prompt, model=model)
                    draft_questions = _extract_questions_from_json(draft)
                except MedtronicGPTError as exc:
                    error = str(exc)
            if not error and draft:
                try:
                    design_update_prompt = build_design_update_prompt(
                        template_text or "",
                        examples,
                        code_context,
                        prior_json=draft,
                        plan_context=plan_text,
                        release_type=release_type,
                        user_instructions=user_instructions,
                    )
                    with _timed(timings, "design_refine_completion_ms"):
                        refined = client.generate_completion(design_update_prompt, model=model)
                    draft = refined
                    draft_questions = _extract_questions_from_json(draft)
                except MedtronicGPTError as exc:
                    if not error:
                        error = f"Design refinement failed: {exc}"
        if action in {"chat", "answers"} and not draft and draft_json_from_form.strip():
            draft = draft_json_from_form
            draft_questions = _extract_questions_from_json(draft)
        if draft:
            draft_json_from_form = draft
        if stored_template:
            final_templates = _dedupe_by_name([stored_template] + final_templates)
            selected_template_name = stored_template.name
        existing_examples: List[StoredFile] = list(stored_inputs.examples)
        existing_names = {ex.name for ex in existing_examples}
        for ex in stored_examples:
            if ex.name not in existing_names:
                existing_examples.append(ex)
                existing_names.add(ex.name)
        persisted_inputs = SavedInputs(
            templates=final_templates,
            examples=_dedupe_by_name(existing_examples),
        )
        if not selected_template_name and persisted_inputs.templates:
            selected_template_name = persisted_inputs.templates[0].name
        save_inputs(persisted_inputs)
        if client and client.last_refresh and remember_credentials:
            stored = StoredCredentials(
                subscription_key=client.subscription_key,
                api_token=client.api_token,
                refresh_token=client.refresh_token,
                api_version=client.api_version,
                base_url=client.base_url,
                path_template=client.path_template,
                model=model,
            )
            save_credentials(stored)
        coverage_note = None
        missing_placeholders: List[str] = []
        coverage_source = draft or draft_json_from_form
        if template_text and coverage_source:
            missing_placeholders = _compute_missing_placeholders(template_text, coverage_source)
            if (
                missing_placeholders
                and client
                and action in {"build", "refine", "answers"}
                and not error
            ):
                update_prompt = build_update_prompt(
                    template_text or "",
                    examples,
                    code_context,
                    coverage_source,
                    answered=answered if "answered" in locals() else [],
                    plan_context=plan_text,
                    release_type=release_type,
                    missing_tokens=missing_placeholders,
                    user_instructions=user_instructions,
                )
                try:
                  with _timed(timings, "refine_completion_ms"):
                    draft = client.generate_completion(update_prompt, model=model)
                    draft_json_from_form = draft
                    draft_questions = _extract_questions_from_json(draft)
                    coverage_source = draft
                    missing_placeholders = _compute_missing_placeholders(template_text, coverage_source)
                except MedtronicGPTError as exc:
                    error = error or str(exc)
            elif missing_placeholders and not client:
                coverage_note = "Provide MedtronicGPT credentials to auto-fill the remaining placeholders."
        if (
          not error
          and client
          and template_text
          and draft  # use the latest JSON after coverage/update prompts
          and action in {"build", "refine", "answers"}
      ):
          with _timed(timings, "functional_requirements_enrichment_ms"):
            draft, fr_raw_latest = _apply_functional_requirements_enrichment(
                draft=draft,
                template_text=template_text,
                prompt=prompt,
                code_context=code_context_text,
                plan_context=plan_text,
                client=client,
                model=model,
            )
          if fr_raw_latest:
              fr_raw = fr_raw_latest
          with _timed(timings, "testing_documentation_alignment_ms"):
            draft, td_raw_latest = _apply_testing_documentation_alignment_enrichment(
                draft=draft,
                template_text=template_text,
                client=client,
                model=model,
            )
          if td_raw_latest:
              fr_raw = td_raw_latest
    try:
        if request.method == "POST":
            final_action = request.form.get("action", "build")
            if final_action in {"build", "refine", "chat", "answers"}:
                run_success = (error is None or str(error).strip() == "")
                user_name = request.form.get("user_name", "").strip() or "unknown"
                log_event(
                    "ui_run",
                    user_name=user_name,
                    app_version=APP_VERSION,
                    action=final_action,
                    model=(request.form.get("model", "").strip() or defaults.get("model") or ""),
                    success=run_success,
                    error=str(error) if error else "",
                    payload={
                        "run_id": run_id,
                        "started_ms": run_started_ms,
                        "release_type": release_type,
                        "template_name": selected_template_name,
                        "has_plan": bool((plan_text or "").strip()),
                        "has_draft": bool((draft_json_from_form or "").strip()),
                        "draft_questions_count": len(draft_questions or []),
                        "missing_placeholders_count": len(missing_placeholders or []),
                        "examples_count": len(persisted_inputs.examples) if persisted_inputs else None,
                        "timings_ms": timings,
                    },
                )
    except Exception:
        app.logger.exception("ui_run telemetry failed")
    return render_template_string(
        TEMPLATE,
        prompt=prompt,
        draft=draft,
        app_version=APP_VERSION,
        error=error,
        defaults=defaults,
        history=history,
        stored=stored,
        saved_inputs=persisted_inputs,
        user_instructions=user_instructions,
        plan_text=plan_text,
        draft_questions=draft_questions,
        code_context=code_context_text,
        code_context_old=code_context_old_text,
        code_context_new=code_context_new_text,
        draft_json=draft_json_from_form,
        coverage_note=coverage_note,
        template_text=template_text,
        missing_placeholders=missing_placeholders,
        release_type=release_type,
        selected_template_name=selected_template_name,
        design_text=design_text,
        fr_raw=fr_raw,
    )
TEMPLATE = """
<!doctype html>
<html lang=\"en\">
<head>
  <meta charset=\"utf-8\">
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">
  <title>Medtronic Validation Draft Builder</title>
  <style>
    :root {
      --bg: #f5f7fb;
      --card: #ffffff;
      --muted: #475569;
      --text: #0f172a;
      --accent: #2563eb;
      --accent-2: #22d3ee;
      --border: rgba(15, 23, 42, 0.08);
      --shadow: 0 24px 70px rgba(15, 23, 42, 0.08);
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: 'Inter', system-ui, -apple-system, sans-serif;
      background: radial-gradient(circle at 12% 10%, rgba(34, 211, 238, 0.18), transparent 26%),
                  radial-gradient(circle at 82% 0%, rgba(37, 99, 235, 0.14), transparent 23%),
                  var(--bg);
      color: var(--text);
      min-height: 100vh;
    }
    a { color: var(--accent); }
    h1, h2, h3 { margin: 0; }
    .page { max-width: 1200px; margin: 0 auto; padding: 32px 24px 48px; }
    .header { display: flex; align-items: center; justify-content: space-between; gap: 16px; margin-bottom: 20px; }
    .badge { padding: 8px 12px; border-radius: 999px; background: linear-gradient(120deg, rgba(34, 211, 238, 0.15), rgba(37, 99, 235, 0.14)); color: var(--accent); font-weight: 600; font-size: 14px; }
    .subtitle { color: var(--muted); margin-top: 10px; line-height: 1.5; }
    .stack { display: flex; flex-direction: column; gap: 12px; }
    .muted { color: var(--muted); }
    .card {
      background: var(--card);
      border-radius: 14px;
      border: 1px solid var(--border);
      padding: 18px;
      box-shadow: 0 14px 40px rgba(15, 23, 42, 0.06);
    }
    .section-body { display: grid; gap: 14px; }
    .panel {
      background: var(--card);
      border-radius: 14px;
      border: 1px solid var(--border);
      padding: 16px 18px;
      box-shadow: 0 14px 40px rgba(15, 23, 42, 0.06);
      position: relative;
    }
    .panel::before {
      content: attr(data-step);
      position: absolute;
      left: 14px;
      top: -12px;
      background: linear-gradient(135deg, #2563eb, #1d4ed8);
      color: #fff;
      padding: 6px 10px;
      border-radius: 12px;
      font-weight: 700;
      font-size: 13px;
      box-shadow: 0 10px 25px rgba(37, 99, 235, 0.35);
    }
    .panel h3 { margin: 6px 0 8px; font-size: 17px; }
    .panel p { margin: 0 0 10px; color: var(--muted); }
    .tag-row { display: flex; align-items: center; gap: 8px; }
    .pill-old { background: rgba(59,130,246,0.1); color: #2563eb; border-color: rgba(59,130,246,0.35); }
    .pill-new { background: rgba(16,185,129,0.1); color: #059669; border-color: rgba(16,185,129,0.3); }
    .card h3 { margin-bottom: 10px; }
    .card p { color: var(--muted); margin: 6px 0 12px; }
    .section { margin-top: 22px; }
    .section-head { display: flex; flex-wrap: wrap; align-items: baseline; gap: 10px; margin-bottom: 8px; }
    .section-head h2 { font-size: 22px; margin: 0; }
    .section-head p { margin: 0; color: var(--muted); }
    .input, textarea, select { width: 100%; padding: 12px 14px; border-radius: 12px; border: 1px solid var(--border); background: #f8fafc; color: var(--text); font-size: 14px; }
    .input:focus, textarea:focus, select:focus { outline: 2px solid rgba(37,99,235,0.3); border-color: rgba(37,99,235,0.35); box-shadow: 0 0 0 3px rgba(37,99,235,0.08); }
    textarea { min-height: 110px; resize: vertical; }
    .checkbox { display: flex; align-items: center; gap: 8px; color: var(--text); }
    .checkbox input { width: 16px; height: 16px; accent-color: var(--accent); }
    .actions { display: flex; flex-wrap: wrap; gap: 12px; margin-top: 14px; align-items: center; }
    .btn {
      border: none; cursor: pointer; border-radius: 12px; padding: 12px 16px; font-weight: 700; font-size: 15px;
      transition: all 0.15s ease; display: inline-flex; align-items: center; gap: 8px;
    }
    .btn-primary { background: linear-gradient(135deg, #2563eb, #1d4ed8); color: #fff; box-shadow: 0 12px 30px rgba(37, 99, 235, 0.25); }
    .btn-ghost { background: #eef2ff; color: #1e293b; border: 1px solid rgba(37, 99, 235, 0.2); }
    .btn:hover { transform: translateY(-1px); }
    .pill { display: inline-flex; align-items: center; gap: 6px; padding: 6px 10px; border-radius: 999px; background: #eef2ff; border: 1px solid rgba(37,99,235,0.18); color: var(--muted); font-size: 12px; }
    .output { white-space: pre-wrap; background: #f8fafc; border: 1px solid var(--border); padding: 16px; border-radius: 14px; min-height: 140px; color: var(--text); }
    .chat { margin-top: 6px; }
    .error { border: 1px solid #ef4444; color: #991b1b; background: #fee2e2; padding: 12px 14px; border-radius: 12px; }
    .tagline { display: flex; gap: 10px; align-items: center; flex-wrap: wrap; color: var(--muted); }
    .loading {
      position: fixed;
      inset: 0;
      background: rgba(15, 23, 42, 0.35);
      display: none;
      align-items: center;
      justify-content: center;
      z-index: 50;
      backdrop-filter: blur(2px);
    }
    .loading.visible { display: flex; }
    .loading-card {
      background: #fff;
      padding: 18px 20px;
      border-radius: 16px;
      box-shadow: var(--shadow);
      display: flex;
      align-items: center;
      gap: 12px;
      min-width: 260px;
      border: 1px solid var(--border);
    }
    .scroll-bottom-btn {
      position: fixed;
      right: 24px;
      bottom: 24px;
      z-index: 60;
      border-radius: 999px;
      padding-inline: 18px;
      box-shadow: 0 16px 40px rgba(15, 23, 42, 0.35);
    }
    .spinner {
      width: 24px;
      height: 24px;
      border: 3px solid rgba(37, 99, 235, 0.25);
      border-top-color: #2563eb;
      border-radius: 999px;
      animation: spin 0.8s linear infinite;
    }
    @keyframes spin { to { transform: rotate(360deg); } }
  </style>
</head>
<body>
  <div class=\"page\">
    <div class=\"loading\" id=\"loading\" aria-live=\"polite\" aria-busy=\"true\">
      <div class=\"loading-card\">
        <div class=\"spinner\" role=\"status\" aria-label=\"Loading\"></div>
        <div>
          <div style=\"font-weight: 700; color: var(--text);\">Working on your answers…</div>
          <div style=\"color: var(--muted); font-size: 14px;\">This may take a few minutes.</div>
        </div>
      </div>
    </div>
    <div class=\"header\">
      <div>
        <div class=\"badge\">Medtronic Validation Draft Builder</div>
      </div>
    </div>
    {% if error %}
      <div class=\"error\"><strong>Error:</strong> {{ error }}</div>
    {% endif %}
    <div class=\"section\">
      <div class=\"section-head\">
        <h2>Inputs</h2>
        <p>Step through the essentials: template, release type, examples, context, and connection.</p>
      </div>
      <form id=\"mainForm\" method=\"post\" enctype=\"multipart/form-data\">
      <input type=\"hidden\" name=\"plan_text\" value=\"{{ plan_text }}\">
      <textarea name=\"draft_json\" style=\"display:none;\">{{ draft or draft_json }}</textarea>
      <input type=\"hidden\" name=\"remove_example\" id=\"removeExampleInput\" value=\"\">
      <div class=\"section-body\">
        <div class="panel" data-step="Step 1">
          <h3>Template & release</h3>
          <p>Pick a template (saved or new), then choose the release type.</p>
          <div class="stack">
            <div>
              <label class="muted" style="font-weight:600;">Upload template</label>
              <input class="input" type="file" name="template_file">
            </div>
            {% if saved_inputs.templates %}
              <div class="stack" style="gap: 8px;">
                <label class="muted" style="font-weight:600;">Select template</label>
                <div style="display:flex; gap:8px; align-items:center;">
                  <select class="input" name="selected_template" id="templateSelect">
                    {% for tmpl in saved_inputs.templates %}
                      <option value="{{ tmpl.name }}" {% if tmpl.name == selected_template_name %}selected{% endif %}>{{ tmpl.name }}</option>
                    {% endfor %}
                  </select>
                  <button type="submit" name="remove_template" id="removeTemplateButton" value="{{ selected_template_name or saved_inputs.templates[0].name }}" class="btn btn-ghost" style="padding:10px 12px;" title="Remove selected template">&#8722;</button>
                </div>
              </div>
            {% endif %}
            <div style="margin-top:10px;">
              <label class="muted" style="font-weight:600;">Template-specific instructions (optional)</label>
              <textarea
                name="user_instructions"
                placeholder="Examples:
            - Ignore sections/tables: ...
            - Do not fill names/signatures; leave blank.
            - Use project codename 'X' everywhere.
            - Treat any bracketed text as instructions to delete, not fill.
            - Ignore example purpose statements in blue; replace with our own."
                style="min-height: 90px;"
              >{{ user_instructions or '' }}</textarea>
              <p class="muted" style="margin:6px 0 0; font-size:12px;">
                These instructions are applied before initial/update logic and should be treated as constraints unless they conflict with the template.
              </p>
            </div>
            <div style="display:grid; gap:10px;">
              <label class="checkbox">
                <input type="radio" name="release_type" value="initial" {% if release_type != 'update' %}checked{% endif %}>
                <span>Initial release (default)</span>
              </label>
              <label class="checkbox" style="align-items:flex-start;">
                <input type="radio" name="release_type" value="update" {% if release_type == 'update' %}checked{% endif %}>
                <span>
                  Update/change: use the last 2–3 validation/quality assurance documents as examples
                  and upload prior + current code/files so deltas are clear.
                </span>
              </label>
              <p class="muted" style="margin:0;">When set to update, extra slots appear for previous vs updated code/context so GPT knows what changed.</p>
            </div>
          </div>
        </div>
        <div class="panel" data-step="Step 2">
          <h3>Examples</h3>
          <p>
            Provide example docs to guide tone and structure &mdash;
            ideally the last 2–3 validation/quality assurance documents for this tool (or a similar one),
            especially when doing an update/change.
          </p>
          <div class="stack">
            <input class="input" type="file" name="examples" multiple>
            {% if saved_inputs.examples %}
              <div class="stack" style="gap:8px;">
                <div style="display:flex; align-items:center; justify-content:space-between; gap:8px;">
                  <div class="pill" style="background: rgba(34,211,238,0.1); color: #067bc7; border-color: rgba(34,211,238,0.25);">
                    Saved examples
                  </div>
                  <button
                    type="button"
                    id="uncheckAllExamples"
                    class="btn btn-ghost"
                    style="padding:6px 10px; font-size:12px;"
                    title="Uncheck all saved examples for this run"
                  >
                    Uncheck all
                  </button>
                </div>
                <div class="stack" style="gap:6px;">
                  {% for example in saved_inputs.examples %}
                    <div style="display:flex; align-items:center; gap:10px;">
                      <label class="checkbox" style="margin:0; flex:1;">
                        <input type="checkbox" name="keep_example_{{ loop.index0 }}" id="keep_example_{{ loop.index0 }}" checked>
                        <span>{{ example.name }}</span>
                      </label>
                      <button type="button" class="btn btn-ghost" data-remove-example="{{ example.name }}" title="Remove example" style="padding:8px 10px;">&#8722;</button>
                    </div>
                  {% endfor %}
                </div>
                <p class="muted" style="margin-top:6px;">
                  Select which saved examples to use for this run. “Uncheck all” keeps them saved but excludes them from the next build.
                </p>
              </div>
            {% endif %}
          </div>
        </div>
        <div class=\"panel\" data-step=\"Step 3\">
          <h3>Code & context</h3>
          <p>Attach relevant code or notes so answers stay anchored to your build.</p>
          <div class="stack">
            <div class="initial-only">
              <div class="tag-row" style="margin-bottom:4px;">
                <span class="pill">Current / general</span>
                <span class="muted">Files and snippets that describe the current build overall.</span>
              </div>
              <div class="file-picker-row">
                <label class="btn btn-ghost" style="padding:8px 12px;">Add files
                  <input type="file" id="codeFilesPicker" multiple style="display:none;">
                </label>
                <label class="btn btn-ghost" style="padding:8px 12px;">Add folder
                  <input type="file" id="codeFolderPicker" webkitdirectory directory multiple style="display:none;">
                </label>
              </div>
              <input class="input" type="file" name="code_files" id="codeFilesManaged" multiple style="display:none;">
              <div id="codeFileList" class="stack" style="gap:6px;"></div>
              <div style="margin-top:6px;">
                <label class="muted" style="font-weight:600;">Provide supporting snippets (current/general)</label>
                <textarea name="code_context" placeholder="Paste notes, links, or code snippets for the current implementation.">{{ code_context }}</textarea>
              </div>
            </div>
            <div class="update-only" style="margin-top: 10px; display:none;">
              <div class="tag-row">
                <span class="pill pill-old">OLD</span>
                <span class="muted" style="margin:0;">Previous code or config</span>
              </div>
              <div class="file-picker-row" style="margin-top:6px;">
                <label class="btn btn-ghost" style="padding:8px 12px;">Add files
                  <input type="file" id="codeFilesOldPicker" multiple style="display:none;">
                </label>
                <label class="btn btn-ghost" style="padding:8px 12px;">Add folder
                  <input type="file" id="codeFolderOldPicker" webkitdirectory directory multiple style="display:none;">
                </label>
              </div>
              <input class="input" type="file" name="code_files_old" id="codeFilesOldManaged" multiple style="display:none;">
              <div id="codeFileOldList" class="stack" style="gap:6px;"></div>
              <div style="margin-top:6px;">
                <label class="muted" style="font-weight:600;">Provide supporting snippets (previous)</label>
                <textarea name="code_context_old" placeholder="Notes or code snippets that describe the PRIOR implementation.">{{ code_context_old }}</textarea>
              </div>
              <div class="tag-row" style="margin-top: 14px;">
                <span class="pill pill-new">NEW</span>
                <span class="muted" style="margin:0;">Updated code or config</span>
              </div>
              <div class="file-picker-row" style="margin-top:6px;">
                <label class="btn btn-ghost" style="padding:8px 12px;">Add files
                  <input type="file" id="codeFilesNewPicker" multiple style="display:none;">
                </label>
                <label class="btn btn-ghost" style="padding:8px 12px;">Add folder
                  <input type="file" id="codeFolderNewPicker" webkitdirectory directory multiple style="display:none;">
                </label>
              </div>
              <input class="input" type="file" name="code_files_new" id="codeFilesNewManaged" multiple style="display:none;">
              <div id="codeFileNewList" class="stack" style="gap:6px;"></div>
              <div style="margin-top:6px;">
                <label class="muted" style="font-weight:600;">Provide supporting snippets (updated)</label>
                <textarea name="code_context_new" placeholder="Notes or code snippets that describe the UPDATED implementation.">{{ code_context_new }}</textarea>
              </div>
            </div>
          </div>
        </div>
        <div class=\"panel\" data-step=\"Step 4\" id=\"connection-card\">
            <div style=\"display:flex; align-items:center; justify-content:space-between; gap:10px;\">
              <div>
                <h3 style=\"margin:6px 0 2px;\">MedtronicGPT connection</h3>
              </div>
            <button type=\"button\" id=\"toggle-connection\" class=\"btn btn-ghost\" style=\"padding:8px 10px;\">Hide</button>
          </div>
          <div id=\"connection-body\" style=\"margin-top: 12px;\">
            <label class=\"checkbox\">
              <input type=\"checkbox\" name=\"use_model\" checked>
              <span>Generate answers with MedtronicGPT</span>
            </label>
            <div style=\"margin-top: 10px; display:grid; gap:10px;\">
              <div class="field">
                <label for="base_url" class="muted" style="font-weight:600;">Base URL</label>
                <select class="input" id="base_url" name="base_url">
                  <option value="https://api.gpt.medtronic.com"
                    {% if defaults.base_url == "https://api.gpt.medtronic.com" %}selected{% endif %}>
                    https://api.gpt.medtronic.com
                  </option>
                  <option value="https://api.gpt-dev.medtronic.com"
                    {% if defaults.base_url == "https://api.gpt-dev.medtronic.com" %}selected{% endif %}>
                    https://api.gpt-dev.medtronic.com
                  </option>
                </select>
              </div>
              <div class=\"field\">
                <label for=\"path_template\" class=\"muted\" style=\"font-weight:600;\">Path template</label>
                <input class=\"input\" id=\"path_template\" type=\"text\" name=\"path_template\" placeholder=\"Path template\" value=\"{{ defaults.path_template }}\">
              </div>
              <div class=\"field\">
                <label for=\"api_version\" class=\"muted\" style=\"font-weight:600;\">API version</label>
                <input class=\"input\" id=\"api_version\" type=\"text\" name=\"api_version\" placeholder=\"API version\" value=\"{{ defaults.api_version }}\">
              </div>
              <div class="field">
                <label for="model" class="muted" style="font-weight:600;">Model</label>
                <select class="input" id="model" name="model">
                  {% set current_model = defaults.model or 'gpt-41' %}
                  <option value="gpt-5" {% if current_model == 'gpt-5' %}selected{% endif %}>gpt-5</option>
                  <option value="gpt-5-mini" {% if current_model == 'gpt-5-mini' %}selected{% endif %}>gpt-5-mini</option>
                  <option value="gpt-5-nano" {% if current_model == 'gpt-5-nano' %}selected{% endif %}>gpt-5-nano</option>
                  <option value="gpt-41" {% if current_model == 'gpt-41' %}selected{% endif %}>gpt-41</option>
                  <option value="gpt-41-mini" {% if current_model == 'gpt-41-mini' %}selected{% endif %}>gpt-41-mini</option>
                  <option value="gpt-41-nano" {% if current_model == 'gpt-41-nano' %}selected{% endif %}>gpt-41-nano</option>
                  <option value="o4-mini" {% if current_model == 'o4-mini' %}selected{% endif %}>o4-mini</option>
                  <option value="o3" {% if current_model == 'o3' %}selected{% endif %}>o3</option>
                  <option value="o3-mini" {% if current_model == 'o3-mini' %}selected{% endif %}>o3-mini</option>
                  <option value="gpt-4o-mini" {% if current_model == 'gpt-4o-mini' %}selected{% endif %}>gpt-4o-mini</option>
                  <option value="gpt-4o" {% if current_model == 'gpt-4o' %}selected{% endif %}>gpt-4o</option>
                  <option value="anthropic.claude-3-5-sonnet-20241022-v2:0" {% if current_model == 'anthropic.claude-3-5-sonnet-20241022-v2:0' %}selected{% endif %}>anthropic.claude-3-5-sonnet-20241022-v2:0</option>
                  <option value="claude-sonnet-3-7" {% if current_model == 'claude-sonnet-3-7' %}selected{% endif %}>claude-sonnet-3-7</option>
                  <option value="claude-sonnet-4" {% if current_model == 'claude-sonnet-4' %}selected{% endif %}>claude-sonnet-4</option>
                  <option value="deepseek-r1" {% if current_model == 'deepseek-r1' %}selected{% endif %}>deepseek-r1</option>
                  <option value="pixtral-large-2502" {% if current_model == 'pixtral-large-2502' %}selected{% endif %}>pixtral-large-2502</option>
                  <option value="llama-maverick-17b-instruct" {% if current_model == 'llama-maverick-17b-instruct' %}selected{% endif %}>llama-maverick-17b-instruct</option>
                </select>
              </div>
              <div class=\"field\">
                <label for=\"subscription_key\" class=\"muted\" style=\"font-weight:600;\">Subscription key</label>
                <input class=\"input\" id=\"subscription_key\" type=\"text\" name=\"subscription_key\" placeholder=\"Subscription key\" value=\"{{ stored.subscription_key or '' }}\">
              </div>
              <div class=\"field\">
                <label for=\"api_token\" class=\"muted\" style=\"font-weight:600;\">API token</label>
                <input class=\"input\" id=\"api_token\" type=\"text\" name=\"api_token\" placeholder=\"API token\" value=\"{{ stored.api_token or '' }}\">
              </div>
              <div class=\"field\">
                <label for=\"refresh_token\" class=\"muted\" style=\"font-weight:600;\">Refresh token</label>
                <input class=\"input\" id=\"refresh_token\" type=\"text\" name=\"refresh_token\" placeholder=\"Refresh token\" value=\"{{ stored.refresh_token or '' }}\">
              </div>
            </div>
            <label class=\"checkbox\" style=\"margin-top: 10px;\">
              <input type=\"checkbox\" name=\"remember_credentials\" checked>
              <span>Remember credentials on this device</span>
            </label>
          </div>
        </div>
        </div>
        <div class="actions" style="margin-top: 12px; justify-content:flex-start;">
          <button class="btn btn-primary" type="submit" name="action" value="build">Generate answers</button>
          <button class="btn btn-ghost" type="submit" name="action" value="clear">
            Clear current run
          </button>
          <label class="checkbox" style="gap:6px;">
            <input type="checkbox" name="remember_inputs" checked>
            <span>Remember uploaded template and examples on this device</span>
          </label>
        </div>
      {% if draft_questions %}
        <div class="card" style="margin-top: 10px;">
          <div class="tagline"><span class="pill">Questions to answer</span><span>Fill these in to update the JSON</span></div>
          <form method="post" id="answersForm">
            <textarea name="draft_json" style="display:none;">{{ draft }}</textarea>
            <input type="hidden" name="user_name" id="userNameHiddenChat" value="">
            <textarea name="user_instructions" style="display:none;">{{ user_instructions or '' }}</textarea>
            <input type="hidden" name="plan_text" value="{{ plan_text }}">
            <textarea name="code_context" style="display:none;">{{ code_context }}</textarea>
            <textarea name="code_context_old" style="display:none;">{{ code_context_old }}</textarea>
            <textarea name="code_context_new" style="display:none;">{{ code_context_new }}</textarea>
            <input type="hidden" name="use_model" value="on">
            <input type="hidden" name="model" id="answersModel" value="{{ defaults.model }}">
            <input type="hidden" name="base_url" value="{{ defaults.base_url }}">
            <input type="hidden" name="api_version" value="{{ defaults.api_version }}">
            <input type="hidden" name="path_template" value="{{ defaults.path_template }}">
            <input type="hidden" name="subscription_key" value="{{ stored.subscription_key }}">
            <input type="hidden" name="api_token" value="{{ stored.api_token }}">
            <input type="hidden" name="refresh_token" value="{{ stored.refresh_token }}">
            {% for q in draft_questions %}
              <div style="margin-top: 12px;">
                <div class="pill" style="margin-bottom: 6px; display: inline-flex;">Question {{ loop.index }}</div>
                <div style="margin-bottom: 6px; color: #0f172a;">{{ q }}</div>
                <textarea name="answer_{{ loop.index0 }}" placeholder="Type your answer..." style="min-height: 70px;"></textarea>
                <input type="hidden" name="question_{{ loop.index0 }}" value="{{ q }}">
              </div>
            {% endfor %}
            <div class="actions" style="margin-top: 12px; gap: 10px;">
              <button class="btn btn-ghost" type="submit" name="action" value="answers">Save answers into JSON</button>
              <button class="btn btn-primary" type="submit" name="action" value="refine">Send answers to GPT</button>
            </div>
          </form>
        </div>
      {% endif %}
    {% if draft %}
      <div class="section">
        <div class="section-head">
          <h2>Answers</h2>
          <p>Review and copy all current mappings.</p>
        </div>
        <div class="card" style="margin-top: 10px;">
          <div class="tagline"><span class="pill">Generated Answers</span><span>Copy</span></div>
          <div class="actions" style="margin-top: 8px; gap: 8px;">
            <div class="pill" id="viewToggleJson" style="cursor: pointer;">JSON view</div>
            <div
              class="pill"
              id="viewToggleFriendly"
              style="cursor: pointer; background: rgba(34,197,94,0.1); color: #22c55e; border-color: rgba(34,197,94,0.3);"
            >
              Easy view
            </div>
            <div
              class="pill"
              id="viewToggleMarkdown"
              style="cursor: pointer;"
            >
              Markdown view
            </div>
            <button type="button" class="btn btn-primary" id="copyAll">Copy all</button>
            <button type="button" class="btn btn-ghost" id="showFrRaw" {% if not fr_raw %}style="display:none;"{% endif %}>
              Functional Requirements raw JSON
            </button>
            <button type="button" class="btn btn-ghost" id="exportWord">Export to Word</button>
          </div>
          <div id="jsonView" class="output" style="margin-top: 10px; white-space: pre-wrap;">{{ draft }}</div>
          <div
            id="friendlyView"
            class="output"
            style="
              margin-top: 10px;
              display: none;
              white-space: normal;        /* override pre-wrap from .output */
              max-height: 60vh;
              overflow-y: auto;
            "
          ></div>
          <div
            id="markdownView"
            class="output"
            style="
              margin-top: 10px;
              display: none;
              white-space: pre-wrap;      /* good for markdown formatting */
              max-height: 60vh;
              overflow-y: auto;
              font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, 'Liberation Mono', 'Courier New', monospace;
            "
          ></div>
          <p style="margin: 10px 0 0; color: #475569;">Toggle between the raw JSON and a simplified list of answers. Use Copy all to grab the current JSON.</p>
        </div>
      </div>
    {% endif %}
      {% if template_text and (draft or draft_json) %}
        <div class="section">
          <div class="section-head">
            <h2>Coverage</h2>
          </div>
          <div class="card" style="margin-top: 10px;">
          <div class="tagline"><span class="pill">Coverage check</span><span>Template placeholders</span></div>
          {% if missing_placeholders %}
            <p style="margin: 6px 0 10px; color: #475569;">These placeholders still need answers:</p>
            <ul style="margin: 0; padding-left: 18px; color: #0f172a;">
              {% for token in missing_placeholders %}
                <li>{{ token }}</li>
              {% endfor %}
            </ul>
              {% if coverage_note %}
                <p style="margin: 10px 0 0; color: #ef4444;">{{ coverage_note }}</p>
              {% endif %}
          {% else %}
            <p style="margin: 6px 0 0; color: #0f172a;">All detected placeholders have values based on the current answers.</p>
          {% endif %}
        </div>
        </div>
      {% endif %}
    <div class="section" style="margin-bottom: 12px;">
      <div class="section-head">
        <h2>Ask Clarifying Questions and Refine Answers</h2>
        <p>Chat stays grounded in your uploaded context and current answers.</p>
      </div>
      <div class="card" style="margin-top: 10px;">
        <div class="tagline"><span class="pill">Clarify or refine</span></div>
        <form method="post" class="chat" id="chatForm">
          <textarea name="user_instructions" style="display:none;">{{ user_instructions or '' }}</textarea>
          <label class="checkbox" style="margin-top: 10px;">
            <input type="checkbox" name="chat_update_json" id="chat_update_json">
            <span>Integrate chat output to draft JSON (structured)</span>
          </label>
          <p class="muted" style="margin: 6px 0 0; font-size: 12px;">
            When enabled, chat returns an updated JSON draft.
          </p>
          <input type="hidden" name="action" value="chat">
          <input type="hidden" name="use_model" value="on">
          <input type="hidden" name="model" id="chatModel" value="{{ defaults.model }}">
          <input type="hidden" name="base_url" value="{{ defaults.base_url }}">
          <input type="hidden" name="api_version" value="{{ defaults.api_version }}">
          <input type="hidden" name="path_template" value="{{ defaults.path_template }}">
          <input type="hidden" name="subscription_key" value="{{ stored.subscription_key }}">
          <input type="hidden" name="api_token" value="{{ stored.api_token }}">
          <input type="hidden" name="refresh_token" value="{{ stored.refresh_token }}">
          <input type="hidden" name="plan_text" value="{{ plan_text }}">
          <textarea name="draft_json" style="display:none;">{{ draft or draft_json }}</textarea>
          <input type="hidden" name="user_name" id="userNameHidden" value="">
          <textarea name="code_context" style="display:none;">{{ code_context }}</textarea>
          <textarea name="code_context_old" style="display:none;">{{ code_context_old }}</textarea>
          <textarea name="code_context_new" style="display:none;">{{ code_context_new }}</textarea>
          <input type="hidden" name="history_json" value='{{ history | tojson }}'>
          <textarea name="chat_input" placeholder="Ask a question or request edits..." style="min-height: 80px;"></textarea>
          <div class="actions" style="margin-top: 10px;">
            <button class="btn btn-ghost" type="submit">Send</button>
            <div class="pill">Chat stays aligned to your uploaded context.</div>
          </div>
        </form>
        {% if history %}
          <div class="output" style="margin-top: 12px;">
            {% for message in history %}
              <div style="margin-bottom: 8px;"><strong>{{ message.role|capitalize }}:</strong> {{ message.content }}</div>
            {% endfor %}
          </div>
        {% endif %}
      </div>
    </div>
    <div class="section">
      <div class="section-head">
        <h2>Feedback</h2>
        <p>Send feedback with the current output attached (JSON + Easy + Markdown).</p>
      </div>
      <div class="card">
        <form method="post" id="feedbackForm">
          <input type="hidden" name="action" value="feedback">
          <input type="hidden" name="user_name" id="userNameHiddenFeedback" value="">
          <input type="hidden" name="model" value="{{ defaults.model }}">
          <textarea name="feedback_text" placeholder="What worked? What didn’t? What should change?" style="min-height:90px;"></textarea>
          <textarea name="draft_json" style="display:none;">{{ draft or draft_json }}</textarea>
          <textarea name="easy_view" id="easyViewHidden" style="display:none;"></textarea>
          <textarea name="markdown_view" id="markdownViewHidden" style="display:none;"></textarea>
          <div class="actions">
            <button class="btn btn-primary" type="submit">Send feedback</button>
          </div>
        </form>
      </div>
    </div>
    <div style="margin-top:24px; color:#64748b; font-size:12px; text-align:center;">
      Version {{ app_version }}
    </div>
    </div>
    <button
      type="button"
      id="scrollBottomBtn"
      class="btn btn-primary scroll-bottom-btn"
      title="Go to bottom"
    >
      ↓ Bottom
    </button>
    {% if fr_raw %}
    <div
      id="frRawModal"
      style="
        position: fixed;
        inset: 0;
        display: none;
        align-items: center;
        justify-content: center;
        background: rgba(15,23,42,0.45);
        z-index: 70;
        backdrop-filter: blur(3px);
      "
    >
      <div
        style="
          background: #ffffff;
          border-radius: 14px;
          border: 1px solid rgba(15,23,42,0.12);
          max-width: 720px;
          width: 90%;
          max-height: 70vh;
          display: flex;
          flex-direction: column;
          box-shadow: 0 24px 70px rgba(15,23,42,0.6);
        "
      >
        <div
          style="
            padding: 10px 14px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            border-bottom: 1px solid rgba(148,163,184,0.35);
            background: linear-gradient(120deg, rgba(37,99,235,0.08), rgba(15,23,42,0.02));
          "
        >
          <div style="font-weight: 600; color:#0f172a; font-size:14px;">Functional requirements – raw model response</div>
          <button
            type="button"
            id="frRawClose"
            class="btn btn-ghost"
            style="padding:6px 10px; font-size:13px;"
          >
            ✕
          </button>
        </div>
        <div style="padding: 10px 14px;">
          <pre
            style="
              margin: 0;
              white-space: pre-wrap;
              word-break: break-word;
              font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, 'Liberation Mono', 'Courier New', monospace;
              font-size: 12px;
              max-height: 50vh;
              overflow-y: auto;
            "
          >{{ fr_raw }}</pre>
        </div>
      </div>
    </div>
    {% endif %}
    <div id="nameModal" style="
      position:fixed; inset:0; display:none; align-items:center; justify-content:center;
      background: rgba(15,23,42,0.45); z-index: 799858697067897; backdrop-filter: blur(3px);
    ">
      <div style="background:#fff; border-radius:14px; padding:16px; width:90%; max-width:420px; border:1px solid rgba(15,23,42,0.12);">
        <div style="font-weight:700; margin-bottom:6px;">Welcome</div>
        <div class="muted" style="margin-bottom:10px;">Enter your full name once (stored on this device) so usage + feedback can be logged.</div>
        <input id="userNameInput" class="input" placeholder="First Last" />
        <div class="actions" style="margin-top:12px; justify-content:flex-end;">
          <button type="button" id="saveUserNameBtn" class="btn btn-primary">Save</button>
        </div>
      </div>
      <form id="exportForm" method="post" action="/export_docx" target="_blank" style="display:none;">
        <input type="hidden" name="view" id="exportView" value="friendly" />
        <input type="hidden" name="open_now" value="1" />
        <textarea name="draft_json" id="exportDraftJson" style="display:none;"></textarea>
        <textarea name="template_text" id="exportTemplateText" style="display:none;">{{ template_text }}</textarea>
        <textarea name="markdown_text" id="exportMarkdownText" style="display:none;"></textarea>
      </form>
    </div>
    <script>
    const loading = document.getElementById('loading');
    const mainForm = document.getElementById('mainForm');
    const answersForm = document.getElementById('answersForm');
    const chatForm = document.getElementById('chatForm');
    const connectionToggle = document.getElementById('toggle-connection');
    const connectionBody = document.getElementById('connection-body');
    const removeExampleInput = document.getElementById('removeExampleInput');
    const removeExampleButtons = Array.from(document.querySelectorAll('[data-remove-example]'));
    const templateSelect = document.getElementById('templateSelect');
    const removeTemplateButton = document.getElementById('removeTemplateButton');
    const releaseValue = '{{ release_type }}';
    const updateSections = Array.from(document.querySelectorAll('.update-only'));
    const initialSections = Array.from(document.querySelectorAll('.initial-only'));
    const releaseRadios = Array.from(document.querySelectorAll('input[name="release_type"]'));
    const rememberInputs = document.querySelector('input[name="remember_inputs"]');
    const templateInput = document.querySelector('input[name="template_file"]');
    const examplesInput = document.querySelector('input[name="examples"]');
    const codeFilesManaged = document.getElementById('codeFilesManaged');
    const templateText = {{ (template_text or "")|tojson }};
    const codeFilesPicker = document.getElementById('codeFilesPicker');
    const codeFolderPicker = document.getElementById('codeFolderPicker');
    const codeFileList = document.getElementById('codeFileList');
    const codeFilesOldManaged = document.getElementById('codeFilesOldManaged');
    const codeFilesOldPicker = document.getElementById('codeFilesOldPicker');
    const codeFolderOldPicker = document.getElementById('codeFolderOldPicker');
    const codeFileOldList = document.getElementById('codeFileOldList');
    const codeFilesNewManaged = document.getElementById('codeFilesNewManaged');
    const codeFilesNewPicker = document.getElementById('codeFilesNewPicker');
    const codeFolderNewPicker = document.getElementById('codeFolderNewPicker');
    const codeFileNewList = document.getElementById('codeFileNewList');
    const uncheckAllExamplesButton = document.getElementById('uncheckAllExamples');
    const scrollBottomBtn = document.getElementById('scrollBottomBtn');
    const modelSelect = document.getElementById('model');
    const answersModelInput = document.getElementById('answersModel');
    const chatModelInput = document.getElementById('chatModel');
    const nameModal = document.getElementById('nameModal');
    const userNameInput = document.getElementById('userNameInput');
    const saveUserNameBtn = document.getElementById('saveUserNameBtn');
    const userNameHidden = document.getElementById('userNameHidden');
    const feedbackForm = document.getElementById('feedbackForm');
    const easyHidden = document.getElementById('easyViewHidden');
    const mdHidden = document.getElementById('markdownViewHidden');
    const exportWordBtn = document.getElementById('exportWord');
    const exportForm = document.getElementById('exportForm');
    const exportView = document.getElementById('exportView');
    const exportDraftJson = document.getElementById('exportDraftJson');
    const exportMarkdownText = document.getElementById('exportMarkdownText');
    if (exportWordBtn && exportForm) {
      exportWordBtn.addEventListener('click', () => {
        const viewToSend = activeView === 'friendly' ? 'friendly' : activeView;
        exportView.value = viewToSend;
        const jsonText = (jsonView && jsonView.textContent) ? jsonView.textContent : rawDraftText();
        exportDraftJson.value = jsonText || '';
        if (viewToSend === 'markdown') {
          exportMarkdownText.value = (markdownView && markdownView.textContent) ? markdownView.textContent : '';
        } else {
          exportMarkdownText.value = '';
        }
        exportForm.submit();
      });
    }
    if (feedbackForm) {
      feedbackForm.addEventListener('submit', () => {
        if (friendlyView && easyHidden) {
          easyHidden.value = friendlyView.innerText || friendlyView.textContent || '';
        }
        if (markdownView && mdHidden) {
          mdHidden.value = markdownView.textContent || '';
        }
        const name = localStorage.getItem('validation_user_name') || '';
        feedbackForm.querySelectorAll('input[name="user_name"]').forEach((el) => el.value = name);
      });
    }
    function syncUserNameToAllForms(name) {
      document.querySelectorAll('input[name="user_name"]').forEach((el) => {
        el.value = name || '';
      });
    }
    function initUserName() {
      const nameModal = document.getElementById('nameModal');
      const existing = localStorage.getItem('validation_user_name') || '';
      if (!existing.trim()) {
        if (nameModal) nameModal.style.display = 'flex';
      } else {
        syncUserNameToAllForms(existing.trim());
      }
    }
    document.addEventListener('DOMContentLoaded', initUserName);
    if (saveUserNameBtn && userNameInput) {
      saveUserNameBtn.addEventListener('click', () => {
        const name = (userNameInput.value || '').trim();
        if (!name) return;
        localStorage.setItem('validation_user_name', name);
        syncUserNameToAllForms(name);
        if (nameModal) nameModal.style.display = 'none';
      });
    }
    function submitWithAction(actionValue) {
      if (!mainForm) return;
      const hidden = document.createElement('input');
      hidden.type = 'hidden';
      hidden.name = 'action';
      hidden.value = actionValue;
      mainForm.appendChild(hidden);
      mainForm.submit();
    }
    let answersReleaseInput = null;
    let chatReleaseInput = null;
    function syncUpdateSections(value) {
      const isUpdate = value === 'update';
      updateSections.forEach((node) => {
        node.style.display = isUpdate ? '' : 'none';
      });
      initialSections.forEach((node) => {
        node.style.display = isUpdate ? 'none' : '';
      });
    }
    function syncReleaseInputs(value) {
      if (answersReleaseInput) answersReleaseInput.value = value;
      if (chatReleaseInput) chatReleaseInput.value = value;
    }
    syncUpdateSections(releaseValue);
    releaseRadios.forEach((radio) => {
      radio.addEventListener('change', (e) => {
        const value = e.target.value;
        syncUpdateSections(value);
        syncReleaseInputs(value);
      });
    });
    if (modelSelect) {
      const syncModelValue = () => {
        const value = modelSelect.value;
        if (answersModelInput) answersModelInput.value = value;
        if (chatModelInput) chatModelInput.value = value;
      };
      syncModelValue();
      modelSelect.addEventListener('change', syncModelValue);
    }
    if (scrollBottomBtn) {
      scrollBottomBtn.addEventListener('click', () => {
        const doc = document.documentElement;
        const height = Math.max(
          doc.scrollHeight,
          document.body.scrollHeight,
          doc.clientHeight
        );
        window.scrollTo({
          top: height,
          behavior: 'smooth',
        });
      });
    }
    if (removeExampleButtons.length && mainForm && removeExampleInput) {
      removeExampleButtons.forEach((btn) => {
        btn.addEventListener('click', (event) => {
          event.preventDefault();
          removeExampleInput.value = btn.getAttribute('data-remove-example') || '';
          if (rememberInputs) rememberInputs.checked = true;
          submitWithAction('remove_example');
        });
      });
    }
    if (templateInput && mainForm) {
      templateInput.addEventListener('change', () => {
        if (rememberInputs) rememberInputs.checked = true;
        submitWithAction('save_template');
      });
    }
    if (examplesInput && mainForm) {
      examplesInput.addEventListener('change', () => {
        if (rememberInputs) rememberInputs.checked = true;
        submitWithAction('save_examples');
      });
    }
    if (uncheckAllExamplesButton) {
      uncheckAllExamplesButton.addEventListener('click', () => {
        const boxes = document.querySelectorAll('input[type="checkbox"][name^="keep_example_"]');
        boxes.forEach((box) => {
          box.checked = false;
        });
      });
    }
    function createManagedFileBucket(pickerEl, folderPickerEl, managedInputEl, listContainerEl) {
      if (!managedInputEl) return null;
      const store = new Map();
      function rebuild() {
        if (typeof DataTransfer === 'undefined') return;
        const dt = new DataTransfer();
        if (listContainerEl) listContainerEl.innerHTML = '';
        store.forEach((file, key) => {
          dt.items.add(file);
          if (listContainerEl) {
            const row = document.createElement('div');
            row.style.display = 'flex';
            row.style.alignItems = 'center';
            row.style.gap = '8px';
            row.style.justifyContent = 'space-between';
            row.style.padding = '6px 10px';
            row.style.border = '1px solid #e2e8f0';
            row.style.borderRadius = '10px';
            row.style.background = 'white';
            const name = document.createElement('span');
            name.textContent = key;
            name.style.flex = '1';
            const removeBtn = document.createElement('button');
            removeBtn.type = 'button';
            removeBtn.className = 'btn btn-ghost';
            removeBtn.textContent = '−';
            removeBtn.style.padding = '6px 10px';
            removeBtn.title = 'Remove file';
            removeBtn.addEventListener('click', () => {
              store.delete(key);
              rebuild();
            });
            row.appendChild(name);
            row.appendChild(removeBtn);
            listContainerEl.appendChild(row);
          }
        });
        managedInputEl.files = dt.files;
      }
      function addFiles(fileList) {
        if (!fileList) return;
        Array.from(fileList).forEach((file) => {
          const key =
            file.webkitRelativePath && file.webkitRelativePath.length
              ? file.webkitRelativePath
              : file.name;
          if (!store.has(key)) {
            store.set(key, file);
          }
        });
        if (rememberInputs) rememberInputs.checked = true;
        rebuild();
      }
      if (pickerEl) {
        pickerEl.addEventListener('change', (event) => {
          addFiles(event.target.files);
          pickerEl.value = '';
        });
      }
      if (folderPickerEl) {
        folderPickerEl.addEventListener('change', (event) => {
          addFiles(event.target.files);
          folderPickerEl.value = '';
        });
      }
      return { addFiles, rebuild, store };
    }
    const currentCodeBucket = createManagedFileBucket(
      codeFilesPicker,
      codeFolderPicker,
      codeFilesManaged,
      codeFileList
    );
    const oldCodeBucket = createManagedFileBucket(
      codeFilesOldPicker,
      codeFolderOldPicker,
      codeFilesOldManaged,
      codeFileOldList
    );
    const newCodeBucket = createManagedFileBucket(
      codeFilesNewPicker,
      codeFolderNewPicker,
      codeFilesNewManaged,
      codeFileNewList
    );
    if (answersForm) {
      const rel = document.createElement('input');
      rel.type = 'hidden';
      rel.name = 'release_type';
      rel.value = releaseValue;
      answersForm.prepend(rel);
      answersReleaseInput = rel;
    }
    if (chatForm) {
      const relChat = document.createElement('input');
      relChat.type = 'hidden';
      relChat.name = 'release_type';
      relChat.value = releaseValue;
      chatForm.prepend(relChat);
      chatReleaseInput = relChat;
    }
    if (templateSelect && removeTemplateButton) {
      const syncRemoveTarget = () => {
        removeTemplateButton.value = templateSelect.value || removeTemplateButton.value;
      };
      syncRemoveTarget();
      templateSelect.addEventListener('change', syncRemoveTarget);
    }
    if (mainForm && loading) {
      mainForm.addEventListener('submit', (event) => {
        const submitter = event.submitter;
        const actionValue = submitter ? submitter.value : mainForm.querySelector('input[name="action"]')?.value;
        if (actionValue === 'build' || actionValue === 'refine') {
          loading.classList.add('visible');
        }
      });
    }
    if (answersForm && loading) {
      answersForm.addEventListener('submit', (event) => {
        const submitter = event.submitter;
        const actionValue = submitter ? submitter.value : answersForm.querySelector('input[name="action"]')?.value;
        if (actionValue === 'refine') {
          loading.classList.add('visible');
        }
      });
    }
    if (chatForm && loading) {
      chatForm.addEventListener('submit', () => {
        loading.classList.add('visible');
      });
    }
    if (connectionToggle && connectionBody) {
      const persisted = localStorage.getItem('medtronic-connection-hidden');
      if (persisted === 'true') {
        connectionBody.style.display = 'none';
        connectionToggle.textContent = 'Show';
      }
      connectionToggle.addEventListener('click', () => {
        const hidden = connectionBody.style.display === 'none';
        const nextHidden = !hidden;
        connectionBody.style.display = nextHidden ? 'none' : '';
        connectionToggle.textContent = nextHidden ? 'Show' : 'Hide';
        localStorage.setItem('medtronic-connection-hidden', String(nextHidden));
      });
    }
    const rawDraft = {{ draft|tojson if draft else 'null' }};
    const jsonView = document.getElementById('jsonView');
    const friendlyView = document.getElementById('friendlyView');
    const toggleJson = document.getElementById('viewToggleJson');
    const markdownView = document.getElementById('markdownView');
    const toggleFriendly = document.getElementById('viewToggleFriendly');
    const toggleMarkdown = document.getElementById('viewToggleMarkdown');
    const copyAll = document.getElementById('copyAll');
    const showFrRawBtn = document.getElementById('showFrRaw');
    const frRawModal = document.getElementById('frRawModal');
    const frRawClose = document.getElementById('frRawClose');
    if (showFrRawBtn && frRawModal) {
      showFrRawBtn.addEventListener('click', () => {
        frRawModal.style.display = 'flex';
      });
    }
    if (frRawModal && frRawClose) {
      frRawClose.addEventListener('click', () => {
        frRawModal.style.display = 'none';
      });
      frRawModal.addEventListener('click', (event) => {
        if (event.target === frRawModal) {
          frRawModal.style.display = 'none';
        }
      });
    }
    function rawDraftText() {
      if (rawDraft === null || rawDraft === undefined) return '';
      if (typeof rawDraft === 'string') return rawDraft;
      try {
        return JSON.stringify(rawDraft, null, 2);
      } catch (e) {
        return '' + rawDraft;
      }
    }
    function parseDraft() {
      if (rawDraft === null || rawDraft === undefined || rawDraft === '') return null;
      if (typeof rawDraft === 'string') {
        try {
          return JSON.parse(rawDraft);
        } catch (e) {
          return null;
        }
      }
      return rawDraft;
    }
    function escapeHtml(str) {
      return String(str)
        .replace(/&/g, '&amp;')
        .replace(/</g, '&lt;')
        .replace(/>/g, '&gt;')
        .replace(/"/g, '&quot;')
        .replace(/'/g, '&#39;');
    }
    function formatValue(value, depth = 0) {
      if (value === null || value === undefined || value === '') return '<span style="color:#94a3b8;">(empty)</span>';
      if (Array.isArray(value)) {
        const items = value
          .map((entry) => `<li>${formatValue(entry, depth + 1)}</li>`)
          .join('');
        return `<ul>${items}</ul>`;
      }
      if (typeof value === 'object') {
        const entries = Object.entries(value)
          .map(([k, v]) => `<li><strong>${escapeHtml(k)}</strong>: ${formatValue(v, depth + 1)}</li>`)
          .join('');
        return `<ul>${entries}</ul>`;
      }
      return escapeHtml(value);
    }
    function hasNonEmptyValue(value) {
      if (value === null || value === undefined) return false;
      if (typeof value === 'string') return value.trim().length > 0;
      if (Array.isArray(value)) {
        return value.some(function (v) { return hasNonEmptyValue(v); });
      }
      if (typeof value === 'object') {
        return Object.keys(value).length > 0;
      }
      return true;
    }
    function orderKeysByTemplate(keys) {
      const original = keys.slice(); // preserves insertion order from Set/Object.keys
      if (!templateText || typeof templateText !== 'string' || !templateText.length) {
        return original;
      }
      const scored = original.map((k, idx) => {
        const pos = templateText.indexOf(k);
        return {
          k,
          idx, // original order for stable tie-break
          pos: pos === -1 ? Number.MAX_SAFE_INTEGER : pos,
        };
      });
      scored.sort((a, b) => {
        if (a.pos !== b.pos) return a.pos - b.pos;
        return a.idx - b.idx;
      });
      return scored.map((x) => x.k);
    }
    function renderFriendly() {
      if (!friendlyView) return;
      const parsed = parseDraft();
      if (!parsed || typeof parsed !== 'object') {
        friendlyView.textContent =
          'Could not parse JSON. Use the JSON view to copy manually.';
        return;
      }
      const placeholdersMap =
        parsed.placeholders &&
        typeof parsed.placeholders === 'object' &&
        !Array.isArray(parsed.placeholders)
          ? parsed.placeholders
          : {};
      const answersList = Array.isArray(parsed.answers) ? parsed.answers : [];
      const questionsList = Array.isArray(parsed.questions) ? parsed.questions : [];
      const answersByPlaceholder = {};
      answersList.forEach((item) => {
        if (!item || typeof item !== 'object') return;
        const ph =
          item.placeholder ||
          item.token ||
          (item.question && String(item.question).trim().length
            ? item.question
            : null);
        if (!ph) return;
        if (!answersByPlaceholder[ph]) answersByPlaceholder[ph] = [];
        answersByPlaceholder[ph].push(item);
      });
      const keySet = new Set();
      Object.keys(placeholdersMap || {}).forEach((k) => keySet.add(k));
      Object.keys(answersByPlaceholder).forEach((k) => keySet.add(k));
      const orderedKeys = orderKeysByTemplate(Array.from(keySet));
      const sections = [];
      const items = [];
      orderedKeys.forEach((ph) => {
        const mapVal =
          placeholdersMap && Object.prototype.hasOwnProperty.call(placeholdersMap, ph)
            ? placeholdersMap[ph]
            : undefined;
        const answerItems = answersByPlaceholder[ph] || [];
        let primaryValue;
        if (hasNonEmptyValue(mapVal)) {
          primaryValue = mapVal;
        }
        for (const ans of answerItems) {
          const v =
            ans.replacement !== undefined ? ans.replacement : ans.answer;
          if (hasNonEmptyValue(v)) {
            if (primaryValue === undefined) {
              primaryValue = v;
            }
          }
        }
        if (!hasNonEmptyValue(primaryValue)) {
          return;
        }
        const valueHtml = formatValue(primaryValue);
        const details = [];
        answerItems.forEach((ans) => {
          const q =
            ans.question && ans.question !== ph
              ? String(ans.question)
              : '';
          const where =
            ans.where || ans.location || '';
          const pieces = [];
          if (q) {
            pieces.push('<em>' + escapeHtml(q) + '</em>');
          }
          if (where) {
            pieces.push(
              '<span style="color:#64748b;">' + escapeHtml(where) + '</span>'
            );
          }
          if (!pieces.length) return;
          details.push('<li>' + pieces.join(' – ') + '</li>');
        });
        const detailsHtml = details.length
          ? '<ul style="margin:4px 0 0 18px;">' + details.join('') + '</ul>'
          : '';
        items.push(`
          <li style="margin-bottom:6px;">
            <div><strong>${escapeHtml(ph)}</strong> → ${valueHtml}</div>
            ${detailsHtml}
          </li>
        `);
      });
      if (items.length) {
        sections.push(
          `<div style="margin-bottom: 10px;">
             <div class="pill" style="margin-bottom:6px;">Placeholders &amp; values</div>
             <ul>${items.join('')}</ul>
           </div>`
        );
      }
      if (questionsList.length) {
        const qList = questionsList
          .map((q) => '<li>' + escapeHtml(q) + '</li>')
          .join('');
        sections.push(
          `<div>
             <div class="pill" style="margin-bottom:6px;">Open questions</div>
             <ul>${qList}</ul>
           </div>`
        );
      }
      friendlyView.innerHTML =
        sections.join('') || 'No parsed placeholders with values available.';
    }
    function scalarToMarkdown(value) {
      if (value === null || value === undefined || value === '') {
        return '_(empty)_';
      }
      return String(value);
    }
    function valueToMarkdown(value, level) {
      level = level || 0;
      var indent = '  '.repeat(level);
      if (value === null || value === undefined || value === '') {
        return '_(empty)_';
      }
      if (Array.isArray(value)) {
        if (!value.length) {
          return '_(empty list)_';
        }
        if (value.every(function (v) {
          return v && typeof v === 'object' && !Array.isArray(v);
        })) {
          return value
            .map(function (obj) {
              var lines = Object.entries(obj).map(function ([k, v]) {
                return indent + '  - **' + k + '**: ' + scalarToMarkdown(v);
              });
              return indent + '-\\n' + lines.join('\\n');
            })
            .join('\\n');
        }
        return value
          .map(function (v) {
            return indent + '- ' + scalarToMarkdown(v);
          })
          .join('\\n');
      }
      if (typeof value === 'object') {
        var parts = Object.entries(value).map(function ([k, v]) {
          return indent + '- **' + k + '**: ' + scalarToMarkdown(v);
        });
        return parts.join('\\n');
      }
      return scalarToMarkdown(value);
    }
        function buildMarkdownFromDraft(parsed) {
      if (!parsed) {
        return '# Answers\\n\\n_No data found in JSON._\\n';
      }
      const placeholdersMap =
        parsed.placeholders &&
        typeof parsed.placeholders === 'object' &&
        !Array.isArray(parsed.placeholders)
          ? parsed.placeholders
          : {};
      const answersList = Array.isArray(parsed.answers) ? parsed.answers : [];
      const questionsList = Array.isArray(parsed.questions) ? parsed.questions : [];
      const answersByPlaceholder = {};
      answersList.forEach(function (item) {
        if (!item || typeof item !== 'object') return;
        const ph =
          item.placeholder ||
          item.token ||
          (item.question && String(item.question).trim().length
            ? item.question
            : null);
        if (!ph) return;
        if (!answersByPlaceholder[ph]) answersByPlaceholder[ph] = [];
        answersByPlaceholder[ph].push(item);
      });
      const keySet = {};
      Object.keys(placeholdersMap || {}).forEach(function (k) {
        keySet[k] = true;
      });
      Object.keys(answersByPlaceholder).forEach(function (k) {
        keySet[k] = true;
      });
      const keys = Object.keys(keySet);
      const orderedKeys = orderKeysByTemplate(keys);
      let md = '# Answers\\n\\n';
      let anyPrinted = false;
      orderedKeys.forEach(function (ph) {
        const mapVal =
          placeholdersMap && Object.prototype.hasOwnProperty.call(placeholdersMap, ph)
            ? placeholdersMap[ph]
            : undefined;
        const answerItems = answersByPlaceholder[ph] || [];
        let primaryValue;
        if (hasNonEmptyValue(mapVal)) {
          primaryValue = mapVal;
        }
        answerItems.forEach(function (ans) {
          const v =
            ans.replacement !== undefined ? ans.replacement : ans.answer;
          if (hasNonEmptyValue(v) && primaryValue === undefined) {
            primaryValue = v;
          }
        });
        if (!hasNonEmptyValue(primaryValue)) {
          return;
        }
        anyPrinted = true;
        md += '## ' + ph + '\\n\\n';
        const valueMarkdown = valueToMarkdown(primaryValue, 0);
        md += (valueMarkdown && valueMarkdown.trim())
          ? valueMarkdown + '\\n\\n'
          : '_(empty)_\\n\\n';  // practically won't hit, because of hasNonEmptyValue
        if (answerItems.length) {
          const detailLines = [];
          answerItems.forEach(function (ans) {
            const q =
              ans.question && ans.question !== ph
                ? String(ans.question)
                : '';
            const where =
              ans.where || ans.location || '';
            if (!q && !where) return;
            let line = '';
            if (q) {
              line += '- _' + q + '_';
            }
            if (where) {
              line += (q ? ' – ' : '- ') + '**Location:** ' + where;
            }
            if (!line) return;
            detailLines.push(line);
          });
          if (detailLines.length) {
            md += 'Additional details:\\n';
            detailLines.forEach(function (line) {
              md += line + '\\n';
            });
            md += '\\n';
          }
        }
      });
      if (!anyPrinted) {
        md += '_No placeholders with values found in JSON._\\n\\n';
      }
      if (questionsList.length) {
        md += '## Remaining questions\\n\\n';
        questionsList.forEach(function (q) {
          md += '- ' + q + '\\n';
        });
        md += '\\n';
      }
      return md;
    }
    function renderMarkdown() {
      if (!markdownView) return;
      const parsed = parseDraft();
      const md = buildMarkdownFromDraft(parsed);
      markdownView.textContent = md;
    }
    let activeView = 'friendly';  // default
    if (toggleJson && toggleFriendly && toggleMarkdown && jsonView && friendlyView && markdownView) {
      function setActiveView(view) {
        activeView = view;
        jsonView.style.display = view === 'json' ? 'block' : 'none';
        friendlyView.style.display = view === 'friendly' ? 'block' : 'none';
        markdownView.style.display = view === 'markdown' ? 'block' : 'none';
        const reset = (el) => {
          if (!el) return;
          el.style.background = 'rgba(148,163,184,0.08)';
          el.style.borderColor = 'rgba(148,163,184,0.4)';
          el.style.color = '#475569';
        };
        reset(toggleJson);
        reset(toggleFriendly);
        reset(toggleMarkdown);
        const activate = (el, bg, border, color) => {
          if (!el) return;
          el.style.background = bg;
          el.style.borderColor = border;
          el.style.color = color;
        };
        if (view === 'json') {
          activate(toggleJson, 'rgba(34,211,238,0.1)', 'rgba(34,211,238,0.3)', '#0369a1');
        } else if (view === 'friendly') {
          activate(toggleFriendly, 'rgba(34,197,94,0.1)', 'rgba(34,197,94,0.3)', '#15803d');
        } else if (view === 'markdown') {
          activate(toggleMarkdown, 'rgba(129,140,248,0.15)', 'rgba(129,140,248,0.4)', '#4338ca');
        }
      }
      toggleJson.addEventListener('click', () => {
        jsonView.textContent = rawDraftText();
        setActiveView('json');
      });
      toggleFriendly.addEventListener('click', () => {
        renderFriendly();
        setActiveView('friendly');
      });
      toggleMarkdown.addEventListener('click', () => {
        renderMarkdown();
        setActiveView('markdown');
      });
      renderFriendly();
      jsonView.textContent = rawDraftText();
      renderMarkdown(); // pre-generate so switching is instant
      setActiveView('friendly');
    }
    if (copyAll) {
      copyAll.addEventListener('click', async () => {
        let toCopy = '';
        if (activeView === 'json') {
          toCopy = jsonView ? jsonView.textContent : rawDraftText();
        } else if (activeView === 'markdown') {
          toCopy = markdownView ? markdownView.textContent : '';
        } else if (activeView === 'friendly') {
          toCopy = friendlyView
            ? (friendlyView.innerText || friendlyView.textContent || '')
            : '';
        } else {
          toCopy = rawDraftText();
        }
        try {
          await navigator.clipboard.writeText(toCopy);
          copyAll.textContent = 'Copied!';
          setTimeout(() => (copyAll.textContent = 'Copy all'), 1200);
        } catch (e) {
          copyAll.textContent = 'Copy failed';
          setTimeout(() => (copyAll.textContent = 'Copy all'), 1200);
        }
      });
    }
  </script>
</body>
</html>
"""
def kill_prior_instances_by_keyword() -> None:
    import os
    import time
    keywords = ("validation-ui", "validationlauncher")
    me_pid = os.getpid()
    parent_pid = os.getppid()
    try:
        import psutil  # type: ignore
    except Exception:
        psutil = None
    if psutil:
        try:
            me = psutil.Process(me_pid)
        except Exception:
            me = None
        exclude = {me_pid, parent_pid}
        if me:
            try:
                exclude.update(p.pid for p in me.parents())
            except Exception:
                pass
            try:
                exclude.update(c.pid for c in me.children(recursive=True))
            except Exception:
                pass
        victims = []
        for p in psutil.process_iter(["pid", "name", "exe", "cmdline"]):
            try:
                pid = p.info["pid"]
                if pid in exclude:
                    continue
                name = (p.info.get("name") or "").lower()
                exe = (p.info.get("exe") or "").lower()
                cmd = " ".join(p.info.get("cmdline") or []).lower()
                haystack = f"{name} {exe} {cmd}"
                if any(k in haystack for k in keywords):
                    victims.append(p)
            except Exception:
                continue
        for p in victims:
            try:
                p.terminate()
            except Exception:
                pass
        try:
            _, alive = psutil.wait_procs(victims, timeout=2.0)
        except Exception:
            alive = victims
        for p in alive:
            try:
                p.kill()
            except Exception:
                pass
        time.sleep(0.2)
        return
    return
if __name__ == "__main__":
    kill_prior_instances_by_keyword()
    import os
    import socket
    import threading
    import time
    import webbrowser
    def _find_open_port(host: str, preferred: int) -> int:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            try:
                s.bind((host, preferred))
                return preferred
            except OSError:
                pass
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            s.bind((host, 0))
            return s.getsockname()[1]
    host = os.getenv("VALIDATION_UI_HOST", "127.0.0.1")
    requested_port = int(os.getenv("VALIDATION_UI_PORT", "8000"))
    port = _find_open_port(host, requested_port)
    def _open_browser() -> None:
        time.sleep(1)
        try:
            webbrowser.open(f"http://{host}:{port}")
        except Exception:
            pass
    threading.Thread(target=_open_browser, daemon=True).start()
    app.run(host=host, port=port, debug=False)